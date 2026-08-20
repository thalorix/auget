#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
 BUFFETT ANALYZER — UN SINGOLO FILE  (v1.5.0)
================================================================================
 40 criteri quantitativi · 12 qualitativi · DCF Owner Earnings · 100% automatico

 NOVITÀ v1.5.0:
   • Bonus "TOTAL": le righe con total/totale vincono sui sotto-voci
     (risolve ricavi=307k preso da "advertising revenue" invece di "Total revenues")
   • norm_line normalizza trattini tipografici (– — ‑ ‐ −) e apici (’ ‘ “ ”):
     ora "Long–term debt" (en-dash SEC) viene riconosciuto
   • Fallback valore-su-riga-successiva esteso a 3 righe
   • LOG diagnostico: se una voce non è trovata ma la label esiste, mostra le
     righe dove la label compare senza valore associato (per calibrazione)

 USO:
   python buffett_analyzer.py                 -> apre la GUI
   python buffett_analyzer.py bilancio.pdf    -> analisi automatica all'avvio

 DIPENDENZE: solo libreria standard (tkinter). PDF: pip install pypdf
================================================================================
"""

import os
import re
import sys
import math
import html
import json
import datetime
import threading
import webbrowser
import traceback
from dataclasses import dataclass
from typing import Optional

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
except ImportError:
    print("ERRORE: tkinter non disponibile. Installa Python con supporto GUI.")
    sys.exit(1)

PdfReader = None
try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

# ============================================================================
# COSTANTI GRAFICHE
# ============================================================================
BG      = "#0b0e14"
PANEL   = "#11151f"
CARD    = "#161c28"
CARD2   = "#1b2333"
BORDER  = "#232c3f"
FG      = "#e8ecf4"
MUTED   = "#8a94ab"
GOLD    = "#f2b632"
GREEN   = "#2ecc71"
AMBER   = "#f39c12"
RED     = "#ff5c5c"
GRAY    = "#5a6478"
BLUE    = "#4da3ff"

F_TITLE  = ("Segoe UI", 20, "bold")
F_SUB    = ("Segoe UI", 9)
F_H2     = ("Segoe UI", 13, "bold")
F_BODY   = ("Segoe UI", 10)
F_SMALL  = ("Segoe UI", 9)
F_TINY   = ("Segoe UI", 8)
F_CARD_T = ("Segoe UI", 10, "bold")
F_MONO   = ("Consolas", 9)

STATUS_META = {
    "pass": (GREEN, "✔ OK"),
    "warn": (AMBER, "◑ PARZIALE"),
    "partial": (AMBER, "◑ PARZIALE"),
    "fail": (RED,   "✘ CRITICO"),
    "nd":   (GRAY,  "… N/D"),
}
QSTATUS_META = {
    "pass":    (GREEN, "✔ SUPERATO"),
    "warn":    (AMBER, "◑ PARZIALE"),
    "partial": (AMBER, "◑ PARZIALE"),
    "fail":    (RED,   "✘ NON SUPERATO"),
    "nd":      (GRAY,  "… N/D"),
}

APP_TITLE = "BUFFETT ANALYZER"
VERSION   = "1.6.0"

# ============================================================================
# SETUP PROGETTO
# ============================================================================
BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
SAMPLES_DIR = os.path.join(BASE_DIR, "samples")
OUTPUT_DIR  = os.path.join(BASE_DIR, "outputs")
SAMPLE_FILE = os.path.join(SAMPLES_DIR, "sample_report.txt")

SAMPLE_REPORT = """AURELIA S.P.A. — RELAZIONE FINANZIARIA ANNUALE 2025
Valori espressi in milioni di Euro

=== CONTO ECONOMICO ===
Ricavi: 2023 4.820, 2024 5.240, 2025 5.710
Costo del venduto: 2023 2.650, 2024 2.860, 2025 3.080
Margine operativo lordo (EBITDA): 1.500
Spese SG&A: 2025 690
Risultato operativo (EBIT): 2023 1.010, 2024 1.150, 2025 1.290
Oneri finanziari (interessi passivi): 45
Utile netto: 2023 720, 2024 830, 2025 940
Ammortamenti: 210
Spese ricerca e sviluppo (R&D): 85
Aliquota fiscale effettiva: 24%

=== STATO PATRIMONIALE ===
Attività correnti: 2.900
Rimanenze di magazzino: 480
Crediti verso clienti: 860
Cassa e equivalenti: 1.150
Totale attivo: 8.400
Passività correnti: 1.350
Debito a lungo termine: 1.600
Debito totale: 2.050
Patrimonio netto: 4.300

=== RENDICONTO FINANZIARIO ===
Flusso di cassa operativo: 1.180
Spese in conto capitale (CapEx): 260
Free cash flow: 920

=== DATI DI MERCATO ===
Azioni in circolazione: 250 milioni
Prezzo azione: 62,40 Euro
Capitalizzazione di mercato: 15.600 milioni
Dividendi distribuiti: 230
"""


def ensure_project():
    os.makedirs(SAMPLES_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if not os.path.exists(SAMPLE_FILE):
        with open(SAMPLE_FILE, "w", encoding="utf-8") as f:
            f.write(SAMPLE_REPORT)


# ============================================================================
# ESTRAZIONE TESTO DA PDF / TXT
# ============================================================================
def _read_text_bytes(path):
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _ocr_pdf(path):
    """OCR opzionale per PDF scannerizzati (pytesseract + pdf2image)."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        return ""
    try:
        return "\n".join(pytesseract.image_to_string(p) for p in convert_from_path(path, dpi=200))
    except Exception:
        return ""


def _read_text_bytes(path):
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _ocr_pdf(path):
    """OCR opzionale per PDF scannerizzati (pytesseract + pdf2image)."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        return ""
    try:
        return "\n".join(pytesseract.image_to_string(p) for p in convert_from_path(path, dpi=200))
    except Exception:
        return ""


def _read_text_bytes(path):
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _ocr_pdf(path):
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        return ""
    try:
        return "\n".join(pytesseract.image_to_string(p) for p in convert_from_path(path, dpi=200))
    except Exception:
        return ""


def collapse_pipe_tables(text):
    out, cur = [], None
    for ln in text.splitlines():
        s = ln.strip()
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            cells = [c for c in cells if c and not re.fullmatch(r"[-: ]+", c)]
            if not cells:
                continue
            if re.search(r"[A-Za-z]", cells[0]):
                if cur is not None:
                    out.append(" ".join(cur))
                cur = list(cells)
            elif cur is not None:
                cur.extend(cells)
            else:
                cur = list(cells)
        else:
            if cur is not None:
                out.append(" ".join(cur)); cur = None
            out.append(ln)
    if cur is not None:
        out.append(" ".join(cur))
    return "\n".join(out)


def extract_text_from_file(path: str) -> str:
    """Loader universale v2: PDF(testo/OCR), HTML, CSV, XLSX, DOCX(paragrafi+TABELLE), TXT/MD."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        if PdfReader is None:
            raise ValueError("Per i PDF serve 'pypdf': pip install pypdf")
        try:
            reader = PdfReader(path)
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception as e:
            raise ValueError(f"Impossibile leggere il PDF: {e}")
        if not text.strip():
            text = _ocr_pdf(path)
            if not text.strip():
                raise ValueError("PDF scannerizzato senza testo e OCR non disponibile.")
        return text

    if ext in (".html", ".htm"):
        raw = _read_text_bytes(path)
        raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
        raw = re.sub(r"(?s)<[^>]+>", " ", raw)
        return html.unescape(re.sub(r"[ \t]+", " ", raw))

    if ext == ".csv":
        import csv, io
        out = []
        for row in csv.reader(io.StringIO(_read_text_bytes(path))):
            cells = [c.strip() for c in row if c and c.strip()]
            if cells:
                out.append("  ".join(cells))
        return "\n".join(out)

    if ext == ".xlsx":
        try:
            import openpyxl
        except ImportError:
            raise ValueError("Per gli XLSX serve 'openpyxl': pip install openpyxl")
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    out.append("  ".join(cells))
        return "\n".join(out)

    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise ValueError("Per i DOCX serve 'python-docx': pip install python-docx")
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                cells = []
                for c in row.cells:
                    txt = c.text.strip()
                    if txt and txt not in cells:
                        cells.append(txt)
                if cells:
                    parts.append("  ".join(cells))
        return "\n".join(parts)

    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    raise ValueError("Formato file non leggibile")



# ============================================================================
# PARSER NUMERICO DI BASE
# ============================================================================
ACCENTS = str.maketrans(
    "àáâäèéêëìíîïòóôöùúûüçÀÁÂÄÈÉÊËÌÍÎÏÒÓÔÖÙÚÛÜÇ",
    "aaaaeeeeiiiioooouuuucAAAAEEEEIIIIOOOOUUUUC",
)

# v1.5: trattini e apici tipografici dei PDF → ASCII
TYPO = {"–": "-", "—": "-", "‐": "-", "‑": "-", "−": "-",
        "’": "'", "‘": "'", "“": '"', "”": '"'}

NUM_RE = re.compile(r"""
      \(\s*[€$]?\s*\d{1,3}(?:[.,]\d{3})+(?:[.,]\d+)?\s*\)     # (1.234,56)
    | [-−]\s*[€$]?\s*\d{1,3}(?:[.,]\d{3})+(?:[.,]\d+)?        # -1.234,5
    | \d{1,3}(?:[.,]\d{3})+(?:[.,]\d+)?                        # 1.234,56
    | \d+(?:[.,]\d{1,2})?\s?%?                                 # 940 / 62,40 / 16,5%
""", re.X)


def parse_number_token(tok: str) -> Optional[float]:
    t = tok.strip()
    if not t:
        return None
    neg = False
    if t.startswith("(") and t.endswith(")"):
        neg, t = True, t[1:-1].strip()
    if t and t[0] in "-−":
        neg, t = True, t[1:].strip()
    t = t.replace("€", "").replace("$", "").replace(" ", "").rstrip("%")
    if not re.search(r"\d", t):
        return None
    dots, commas = t.count("."), t.count(",")
    if dots and commas:
        if t.rfind(".") > t.rfind(","):
            t = t.replace(",", "")
        else:
            t = t.replace(".", "").replace(",", ".")
    elif commas:
        if commas >= 2:
            t = t.replace(",", "")
        else:
            head, _, tail = t.partition(",")
            if tail.isdigit() and len(tail) == 3 and head not in ("", "0"):
                t = t.replace(",", "")
            else:
                t = t.replace(",", ".")
    elif dots:
        if dots >= 2:
            t = t.replace(".", "")
        else:
            head, _, tail = t.partition(".")
            if tail.isdigit() and len(tail) == 3 and head not in ("", "0"):
                t = t.replace(".", "")
    try:
        v = float(t)
    except ValueError:
        return None
    return -v if neg else v


def norm_line(line: str) -> str:
    for a, b in TYPO.items():
        line = line.replace(a, b)
    return line.lower().translate(ACCENTS).replace("\u00a0", " ")


def collapse_pipe_tables(text):
    """v1.8: compatta tabelle markdown/DOCX (una cella per riga) in righe logiche."""
    out = []
    cur = None
    for ln in text.splitlines():
        s = ln.strip()
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            cells = [c for c in cells if c and not re.fullmatch(r"[-: ]+", c)]
            if not cells:
                continue
            if re.search(r"[A-Za-z]", cells[0]):
                if cur is not None:
                    out.append(" ".join(cur))
                cur = list(cells)
            elif cur is not None:
                cur.extend(cells)
            else:
                cur = list(cells)
        else:
            if cur is not None:
                out.append(" ".join(cur))
                cur = None
            out.append(ln)
    if cur is not None:
        out.append(" ".join(cur))
    return "\n".join(out)


def strip_years(line: str) -> str:
    years = re.findall(r"\b(?:19|20)\d{2}\b", line)
    if len(years) >= 2:
        return re.sub(r"\b(?:19|20)\d{2}\b", " ", line)
    return re.sub(r"\b(?:19|20)\d{2}\b(?=\s*[€$]?\s*\d)", " ", line)


def extract_numbers_from_line(line: str, allow_percent=False):
    clean = strip_years(line)
    out = []
    for m in NUM_RE.finditer(clean):
        tok = m.group(0)
        is_pct = tok.rstrip().endswith("%")
        if is_pct and not allow_percent:
            continue
        v = parse_number_token(tok)
        if v is None:
            continue
        if not is_pct:
            nxt = clean[m.end(): m.end() + 16].lower()
            if re.match(r"\s*(miliard|mld|bn|billion)", nxt):
                v *= 1000.0
            elif re.match(r"\s*(milion|mln|m)", nxt):
                v *= 1.0
        out.append((v, is_pct))
    out = [x for x in out if not (x[1] is False and x[0] == int(x[0]) and 1900 <= x[0] <= 2099)]
    return out


# ============================================================================
# MOTORE DI ESTRAZIONE v2+v5
# ============================================================================
SPECS = [
    ("ricavi",       [r"ricav", r"revenue", r"fatturato", r"\bsales\b", r"turnover",
                      r"valore della produzione", r"total net revenue", r"total revenue"],
                     [r"margine", r"costo", r"\bcost\b", r"cost of", r"crescita", r"growth", r"variazione", r"incidenza"], False),
    ("cogs",         [r"costo del venduto", r"cost of goods sold", r"\bcogs\b", r"cost of sales",
                      r"cost of revenue", r"costo dei ricavi", r"costi della produzione",
                      r"cost of services"], [], False),
    ("ebitda",       [r"ebitda", r"margine operativo lordo"], [r"margine ebitda"], False),
    ("ebit",         [r"ebit(?!da)", r"risultato operativo", r"operating income",
                      r"operating profit", r"utile operativo", r"reddito operativo",
                      r"income from operations"],
                     [r"ebitda", r"margine operativo lordo", r"margine operativo", r"\bros\b"], False),
    ("net_income",   [r"utile netto di pertinenza", r"utile netto del periodo", r"utile netto",
                      r"risultato netto del periodo", r"risultato netto", r"net income",
                      r"profit for the (?:year|period)", r"earnings attributable"],
                     [r"ricav", r"revenue", r"margine", r"margin", r"intermediazione",
                      r"per azione", r"per share", r"attivo", r"assets", r"other comprehensive"], False),
    ("depreciation", [r"ammortament", r"depreciation"], [], False),
    ("capex",        [r"capex", r"conto capitale", r"capital expenditure", r"capital expenditures",
                      r"investimenti in immobilizzazioni", r"acquisto di immobilizzazioni",
                      r"investimenti in attivit (?:materiali|immateriali|fisse)",
                      r"purchases? of property and equipment", r"additions to property and equipment",
                      r"purchase of property and equipment",
                      r"acquisition of property, plant and equipment"],
                     [r",\s*net", r"accumulated"], False),
    ("cfo",          [r"flusso di cassa operativo", r"flusso di cassa da attivit",
                      r"cash flow operativo", r"cash flow from operat", r"operating cash flow",
                      r"flusso di cassa della gestione operativa", r"cassa generata dalla gestione",
                      r"gestione operativa", r"net cash provided by operating activities",
                      r"cash provided by operating activities", r"net cash from operating", r"net cash from operations",
                      r"operating activities"],
                     [r"free", r"libero", r"investiment", r"finanziari"], False),
    ("fcf",          [r"free cash flow", r"flusso di cassa libero", r"\bfcf\b"], [], False),
    ("equity",       [r"totale patrimonio netto", r"patrimonio netto di gruppo", r"patrimonio netto complessivo",
                      r"patrimonio netto", r"stockholders'? equity", r"shareholders'? equity", r"total equity",
                      r"equity attributable", r"capitale proprio", r"net assets"],
                     [r"vincolato", r"noncontrolling", r"non controlling", r"minority", r"redeemable",
                      r"passivit", r"liabilit", r"attivo", r"assets", r"per azione", r"per share",
                      r"other comprehensive"], False),
    ("total_debt",   [r"debito totale", r"total debt", r"debiti totali",
                      r"indebitamento totale", r"debito finanziario lordo",
                      r"total borrowings"], [], False),
    ("net_debt",     [r"posizione finanziaria netta", r"net debt", r"indebitamento netto"], [], False),
    ("lt_debt",      [r"(?:debito|finanziamenti).{0,20}lungo termine", r"long[- ]term debt",
                      r"obbligazioni", r"finanziamenti passivi", r"debt,? less current portion",
                      r"senior notes", r"notes payable", r"term debt"], [r"current portion"], False),
    ("cassa",        [r"cassa e (?:mezzi )?equivalenti", r"cash and (?:cash )?equivalents",
                      r"disponibilit liquide", r"cassa e banche", r"liquidit"], [], False),
    ("total_assets", [r"totale attiv", r"total assets", r"attivo totale",
                      r"totale dell.attivo"], [], False),
    ("current_assets",[r"attivit correnti", r"attivo corrente", r"current assets",
                       r"totale attivit correnti", r"total current assets"], [], False),
    ("current_liab", [r"passivit correnti", r"passivo corrente", r"current liabilities",
                      r"totale passivo corrente", r"passivit a breve", r"total current liabilities"], [], False),
    ("inventory",    [r"rimanenze", r"inventor"], [], False),
    ("crediti",      [r"crediti (?:commerciali|verso clienti)", r"accounts receivable",
                      r"trade receivables", r"crediti commerciali"], [], False),
    ("interessi",    [r"interessi passivi", r"interest expense", r"oneri finanziari"], [r"basis points"], False),
    ("sga",          [r"sg&a", r"\bsga\b", r"spese generali", r"selling,? general",
                      r"administrative expenses", r"costi commerciali",
                      r"spese per servizi", r"costi del personale",
                      r"selling and marketing", r"general and administrative"], [], False),
    ("rnd",          [r"ricerca e sviluppo", r"r&d", r"research and development"], [], False),
    ("shares",       [r"azioni in circolazione", r"shares outstanding", r"common stock outstanding",
                      r"numero (?:di )?azioni", r"number of shares", r"azioni ordinarie",
                      r"shares of common stock", r"outstanding as of",
                      r"common stock.{0,40}outstanding", r"^common stock\b",
                      r"number of outstanding shares", r"outstanding shares of each class"],
                     [r"weighted average", r"media ponderata", r"authorized", r"preferred"], False),
    ("price",        [r"prezzo (?:per )?azione", r"share price", r"stock price",
                      r"prezzo di chiusura", r"closing price", r"market price per share"], [], False),
    ("market_cap",   [r"capitalizzazione", r"market cap"], [], False),
    ("dividendi",    [r"dividendi (?:distribuiti|pagati|totali|deliberati)", r"dividends (?:paid|declared|total)",
                      r"\bdividendi\b", r"\bdividends\b", r"dividends? paid"],
                     [r"payout", r"per azione", r"per share"], False),
    ("tax_rate",     [r"aliquota fiscale", r"tax rate", r"aliquota effettiva",
                      r"imposte sul reddito", r"effective tax rate"], [], True),
    ("goodwill",     [r"goodwill"], [r"impairment"], False),
    ("intangibles",  [r"intangible assets,? net", r"intangibles, net", r"altre attivit immateriali"], [], False),
    ("nim",          [r"net interest margin", r"margine di interesse"], [], True),
    ("npl",          [r"non[- ]performing", r"sofferenze", r"\bnpl\b"], [], True),
    ("combined_ratio", [r"combined ratio"], [], True),
    ("cost_income",  [r"cost/income", r"cost income ratio"], [], True),

    ("rpo",          [r"remaining performance obligation", r"performance obligation.{0,30}remaining",
                      r"backlog", r"ordinativi residui"], [], False),
    ("deferred_rev", [r"deferred revenue", r"unearned revenue", r"ricavi differiti"], [], False),

    ("loans",        [r"crediti verso clientela", r"loans and advances to customers",
                      r"totale crediti", r"gross loans"], [], False),

    ("npl_amount",   [r"sofferenze lorde", r"crediti deteriorati lordi", r"gross non[- ]performing",
                      r"non[- ]performing loans (?:gross|total)"], [], False),

    ("roe_reported", [r"\broe\b", r"return on equity", r"ritorno sul capitale netto"], [r"cost of"], True),
    ("rote",         [r"return on tangible equity", r"\brote\b"], [], True),
    ("roa_reported", [r"\broa\b", r"return on assets"], [], True),

    ("ricavi_bancari", [r"margine di intermediazione", r"net banking income", r"total operating income",
                        r"net interest income", r"ricavi operativi", r"operating income total"],
                     [r"per azione", r"per share"], False),

    ("eps_reported", [r"utile (?:netto )?per azione", r"earnings per share", r"basic earnings per share",
                      r"diluted earnings per share", r"utile per azione base"],
                     [r"per share or rsu", r"dividends"], False),

]

POS_KEYS = ["ricavi", "cogs", "depreciation", "capex", "cassa", "crediti",
            "inventory", "sga", "rnd", "interessi", "current_assets",
            "total_assets", "lt_debt", "total_debt"]


def _apply_custom_patterns():
    """Pattern personalizzati da patterns_custom.json: calibrazione senza modificare il codice."""
    p = os.path.join(BASE_DIR, "patterns_custom.json")
    if not os.path.exists(p):
        return
    try:
        with open(p, encoding="utf-8") as f:
            cust = json.load(f)
    except Exception:
        return
    idx = {k: i for i, (k, _i, _e, _p) in enumerate(SPECS)}
    for e in cust.get("patterns", []):
        key = e.get("key", "")
        inc = list(e.get("include", []))
        exc = list(e.get("exclude", []))
        if not key or not inc:
            continue
        if key in idx:
            k, i2, e2, p2 = SPECS[idx[key]]
            SPECS[idx[key]] = (k, list(i2) + inc, list(e2) + exc, p2)
        else:
            SPECS.append((key, inc, exc, bool(e.get("percent", False))))
            idx[key] = len(SPECS) - 1

_apply_custom_patterns()


def extract_years(line: str):
    return [int(y) for y in re.findall(r"\b(?:19|20)\d{2}\b", line)]


def clean_vals(vals):
    if len(vals) >= 2:
        mx = max(abs(v) for v in vals)
        if mx > 100:
            vals = [v for v in vals if not (v == int(v) and abs(v) <= 20)] or vals
    if len(vals) >= 2:
        mx = max(abs(v) for v in vals)
        if mx > 10000:
            vals = [v for v in vals if abs(v) >= mx / 50] or vals
    if len(vals) >= 3:
        out = []
        for i, v in enumerate(vals):
            others = [w for j, w in enumerate(vals) if j != i]
            is_delta = any(abs(v - (a - b)) < 1e-6 for a in others for b in others if a != b)
            if not is_delta:
                out.append(v)
        if out:
            vals = out
    return vals


def pick_value(vals, years):
    series = list(vals)
    if years and len(years) == len(series):
        if years[0] > years[-1]:
            series.reverse()
        cur = series[-1]
    else:
        cur = series[0]
        series.reverse()  # serie cronologica per CAGR (convenzione SEC/IT: corrente prima)
    return cur, series


def detect_unit(text: str) -> float:
    head = text[:4000]
    if re.search(r"in\s+thousands|in\s+migliaia|in\s+units of thousands", head, re.I):
        return 1e3
    if re.search(r"in\s+millions|in\s+milioni", head, re.I):
        return 1e6
    if re.search(r"in\s+billions|in\s+miliardi", head, re.I):
        return 1e9
    return 1.0


def _vals_from_nums(nums, allow_pct):
    if allow_pct:
        return [v / 100.0 for v, pct in nums if pct and 0 < abs(v) <= 100]
    return clean_vals([v for v, pct in nums if not pct])


def extract_financials(text: str):
    D, S, log = {}, {}, []

    UNIT = detect_unit(text)
    if UNIT != 1.0:
        D["_unit"] = UNIT
        log.append(("warn", f"⚠ Unità rilevata dall'intestazione: valori ×{UNIT:.0f}."))

    if text.count("\n|") > 30:
        text = collapse_pipe_tables(text)
        log.append(("warn", "Tabelle pipe (DOCX/HTML) compattate in righe logiche"))
    lines_raw = [ln.strip() for ln in re.split(r"[\n\r]+", text) if ln.strip()]
    lines_norm = [norm_line(ln) for ln in lines_raw]

    # ---- PASSO 1: candidati multipli + fallback righe successive + bonus TOTAL ----
    cands_by_key = {}
    label_hits = {}
    for key, inc_pats, exc_pats, allow_pct in SPECS:
        cands = []
        for idx, line in enumerate(lines_norm):
            if not re.search(r"\d", line):
                continue
            if any(re.search(p, line) for p in exc_pats):
                continue
            starts = [m.start() for p in inc_pats for m in [re.search(p, line)] if m]
            if not starts:
                continue
            lab_pos = min(starts)
            penalty = 0.0
            vals = _vals_from_nums(extract_numbers_from_line(lines_raw[idx], allow_percent=allow_pct), allow_pct)
            yr_line = lines_raw[idx]
            is_pipe = line.lstrip().startswith("|")
            if not vals:
                for off in range(1, 9 if is_pipe else 5):
                    if idx + off >= len(lines_raw):
                        break
                    vals2 = _vals_from_nums(extract_numbers_from_line(lines_raw[idx + off], allow_percent=allow_pct), allow_pct)
                    if vals2:
                        vals, yr_line, penalty = vals2, lines_raw[idx + off], (1.0 if is_pipe else 1.5)
                        break
            if not vals:
                label_hits.setdefault(key, []).append(idx + 1)
                continue
            years = extract_years(yr_line)
            cur, series = pick_value(vals, years)

            lab_clean = re.sub(r"^[^a-z]+", "", line)
            starts2 = [m.start() for p in inc_pats for m in [re.search(p, lab_clean)] if m]
            lab_pos2 = min(starts2) if starts2 else lab_pos
            score = -penalty
            if lab_pos2 == 0:       score += 3.0
            elif lab_pos2 < 40:     score += 1.5
            if 1 <= len(vals) <= 3: score += 2.0
            elif len(vals) > 5:     score -= 3.0
            mx = max(abs(v) for v in vals)
            if mx >= 100:           score += 1.0
            if mx > 1e12:           score -= 3.0
            if (key in POS_KEYS or key == "shares") and cur < 0:
                score -= 2.0
            if re.search(r"\btotal(?:e)?\b", line):
                score += 2.0
            if key == "shares":
                if "outstanding" in line and "repurchased" not in line and "authorized" not in line:
                    score += 3.0
                elif "repurchased" in line or "issued" in line:
                    score -= 3.0
            if is_pipe or "|" in yr_line:
                score += 2.5
            if len(line) > 110:
                score -= 2.5
            if re.search(r"\b(increased|decreased)\b", line) and re.search(r"\b(billion|million)\b", line):
                score -= 3.0   # frasi "delta" del MD&A: non sono valori di bilancio
            cands.append([score, idx, cur, series])
            if len(cands) >= 400:
                break
        if cands:
            cands_by_key[key] = cands

    # ---- PASSO 2: scala di riferimento dagli anchor ----
    anchors = ["ricavi", "net_income", "equity", "total_assets", "ebit"]
    anchor_logs = []
    for k in anchors:
        if k in cands_by_key:
            best = max(cands_by_key[k], key=lambda c: c[0])
            if best[2] and abs(best[2]) > 1:
                anchor_logs.append(math.log10(abs(best[2])))
    scale = None
    if anchor_logs:
        anchor_logs.sort()
        scale = anchor_logs[len(anchor_logs) // 2]

    # v1.6: bonus plausibilità azioni (un valore tra 1M e 1.000Mld è credibile; "1" no)
    if "shares" in cands_by_key:
        for c in cands_by_key["shares"]:
            if 1e5 <= c[2] <= 1e12:
                c[0] += 3.0

    # ---- PASSO 3: scelta finale con allineamento unità ----
    for key, cands in cands_by_key.items():
        def adjusted(c):
            sc = c[0]
            if key != "shares" and scale is not None and c[2] and abs(c[2]) > 1:
                sc -= 0.8 * abs(math.log10(abs(c[2])) - scale)
            return sc
        best = max(cands, key=adjusted)
        D[key] = best[2]
        S[key] = best[3]
        serie = " → ".join(fmt(v, 0) for v in best[3]) if len(best[3]) > 1 else ""
        log.append(("ok", f"{key:14s} = {fmt(best[2])}  {serie}  "
                          f"(riga {best[1]+1}, score {adjusted(best):.1f}, {len(cands)} candidati)"))
    for spec in SPECS:
        if spec[0] not in D:
            hits = label_hits.get(spec[0])
            if hits:
                log.append(("nd", f"{spec[0]:14s} = label presente (righe {hits[:3]}) ma nessun valore associato"))
            else:
                log.append(("nd", f"{spec[0]:14s} = non trovato nel documento"))

    # ---- SANITÀ: segno coerente per grandezze positive ----
    for k in POS_KEYS:
        if k in D and D[k] < 0:
            D[k] = abs(D[k])
            log.append(("warn", f"{k:14s} = valore negativo (riga di variazione?) → convertito in assoluto"))
    if "dividendi" in D and D["dividendi"] < 0:
        D["dividendi"] = abs(D["dividendi"])
        log.append(("warn", "dividendi     = valore negativo nel rendiconto → convertito in assoluto"))
    if "tax_rate" in D and not (0.05 <= D["tax_rate"] <= 0.70):
        log.append(("warn", f"tax_rate      = scartato valore non plausibile ({fmt(D['tax_rate']*100,1)}%)"))
        del D["tax_rate"]

    for k in ("lt_debt", "total_debt"):
        if k in D and D[k] is not None and 0 < D[k] < 10:
            log.append(("warn", f"{k:14s} = scartato valore immateriale/errato ({fmt(D[k],1)})"))
            del D[k]

    # ---- DERIVAZIONI E STIME ----
    def derive(key, val, how):
        D[key] = val
        log.append(("warn", f"{key:14s} = {how}"))

    if "ebitda" not in D and "ebit" in D and "depreciation" in D:
        derive("ebitda", D["ebit"] + D["depreciation"], "derivato da EBIT + Ammortamenti")
    if "ebit" not in D and "ebitda" in D and "depreciation" in D:
        derive("ebit", D["ebitda"] - D["depreciation"], "derivato da EBITDA − Ammortamenti")
    if "capex" not in D and "cfo" in D and "fcf" in D:
        derive("capex", D["cfo"] - D["fcf"], "derivato da CFO − FCF")
    if "capex" not in D and "depreciation" in D:
        derive("capex", D["depreciation"], "STIMATO ≈ Ammortamenti (proxy CapEx di manutenzione)")
    if "cfo" not in D and "net_income" in D and "depreciation" in D:
        derive("cfo", D["net_income"] + D["depreciation"], "STIMATO ≈ Utile Netto + Ammortamenti (proxy)")
    if "fcf" not in D and "cfo" in D and "capex" in D:
        derive("fcf", D["cfo"] - D["capex"], "derivato da CFO − CapEx")
    if "total_debt" not in D and "net_debt" in D and "cassa" in D:
        derive("total_debt", D["net_debt"] + D["cassa"], "derivato da PFN + Cassa")
    if "total_debt" not in D and "lt_debt" in D:
        derive("total_debt", D["lt_debt"], "derivato da Debito LT (quota breve non trovata)")
    if "lt_debt" not in D and "total_debt" in D:
        D["lt_debt"] = D["total_debt"]
    if "gross_profit" not in D and "ricavi" in D and "cogs" in D:
        D["gross_profit"] = D["ricavi"] - D["cogs"]
    if "cogs" not in D and "ricavi" in D and "gross_profit" in D:
        derive("cogs", D["ricavi"] - D["gross_profit"], "derivato da Ricavi − Margine Lordo")
    if "current_assets" not in D:
        comps = [D.get(k) for k in ("cassa", "crediti", "inventory")]
        present = [c for c in comps if c is not None]
        if len(present) >= 2:
            derive("current_assets", sum(present), "derivato da Cassa + Crediti + Rimanenze")
    if "selling_marketing" in D and "sga" in D and D["sga"] < D["selling_marketing"]:
        D["sga"] = D["sga"] + D["selling_marketing"]
        log.append(("warn", "sga           = sommato Sales&Marketing + General&Administrative"))
    if "sga" not in D and "total_opex" in D and "rnd" in D:
        derive("sga", D["total_opex"] - D["rnd"], "derivato da Total operating expenses - R&D")
    if "tax_rate" not in D and "taxes" in D and "ebt" in D and D["ebt"]:
        tr = D["taxes"] / D["ebt"]
        if 0.05 <= tr <= 0.70:
            derive("tax_rate", tr, "derivato da Imposte / Utile ante imposte")

    # ---- SANITÀ: AZIONI (scala + plausibilità EPS) ----
    # v1.7: rilevatore azienda senza debito ("debt-free", "no outstanding borrowings")
    if re.search(r"debt[-\s]?free|no (?:outstanding )?borrowings|no amounts (?:were )?outstanding under (?:our|the) (?:revolving )?credit facility", text, re.I):
        if "total_debt" not in D:
            D["total_debt"] = 0.0
            log.append(("warn", "total_debt    = 0 (azienda dichiarata debt-free / nessun borrowing)"))
        if "lt_debt" not in D:
            D["lt_debt"] = 0.0
            log.append(("warn", "lt_debt       = 0 (azienda dichiarata debt-free / nessun borrowing)"))

    ni = D.get("net_income")
    sh = D.get("shares")
    if sh is not None:
        U = D.get("_unit", 1.0)
        if sh > 1e5:
            if U == 1.0 and ni:
                for f in (1.0, 1e3, 1e6):
                    if 0.05 <= abs(ni * f / sh) <= 1000:
                        U = f
                        break
                D["_unit"] = U
            sh_doc = sh / U
        else:
            sh_doc = sh
        eps_test = safe_div(ni, sh_doc) if ni else None
        if sh_doc <= 0 or eps_test is None or not (0.05 <= abs(eps_test) <= 1000):
            log.append(("warn", f"shares        = SCARTATE: EPS {fmt(eps_test,0)} non plausibile (riga par value?)"))
            del D["shares"]
        else:
            if sh_doc != sh:
                log.append(("warn", f"shares        = azioni assolute → scala documento (/{U:.0f}) = {fmt(sh_doc,0)}"))
            D["shares"] = sh_doc

    if "price" in D and "shares" in D and ni:
        pe_test = safe_div(D["price"], safe_div(ni, D["shares"]))
        if pe_test is not None and 0 < pe_test < 4:
            log.append(("warn", f"price         = SCARTATO: P/E {fmt(pe_test,1)} non plausibile (riga non di mercato)"))
            del D["price"]

    # ---- SANITÀ: PREZZO non validabile → scartato ----
    if "shares" not in D and "market_cap" not in D and "price" in D:
        log.append(("warn", "price         = SCARTATO: prezzo non presente/validabile nel filing (servono azioni o market cap)"))
        del D["price"]

    # v1.9: istituto finanziario -> metriche corporate di leva/liquidita' non significative
    if re.search(r"\bcet ?1\b|common equity tier|cost/income", text, re.I):
        D["_bank"] = 1
        log.append(("warn", "Istituto finanziario rilevato: leva/liquidita' valutate come N/D (non significative)"))
    _bn = len(re.findall(r"\b(net interest income|margine di intermediazione|loans and advances|crediti verso clientela|deposits? (?:from|della) clientela|cet ?1|capital adequacy|non[- ]performing|sofferenze|proventi di interessi|interest income)\b", text, re.I))
    if _bn >= 2 and not D.get("_bank"):
        D["_bank"] = 1
        log.append(("warn", "Settore finanziario rilevato dai termini tecnici: metriche industriali disattivate"))
    if "market_cap" in D and D.get("price") and D.get("shares"):
        mc_calc = D["price"] * D["shares"]
        if not (0.5 * mc_calc <= D["market_cap"] <= 2.0 * mc_calc):
            log.append(("warn", f"market_cap    = scartato valore incoerente ({fmt(D['market_cap'],0)}) vs prezzo x azioni ({fmt(mc_calc,0)})"))
            del D["market_cap"]
    # v2.0: flag qualitativi dal testo (restatement, opinione auditor)
    D["_restatement"] = 1 if re.search(r"\brestatement\b|\brestated\b", text, re.I) else 0
    D["_unqualified"] = 1 if re.search(r"unqualified opinion", text, re.I) else (0 if re.search(r"audited", text, re.I) else None)
    # v2.0: serie grafico performance azionaria (indice base 100, 6 punti)
    mperf = re.search(r"cumulative total return", text, re.I)
    if mperf:
        for ln in text[mperf.start(): mperf.start() + 4000].splitlines():
            nums = [v for v, pct in extract_numbers_from_line(ln) if not pct]
            if len(nums) == 6 and abs(nums[0] - 100) < 1e-6:
                S["_perf"] = nums
                log.append(("ok", f"perf_index    = {fmt(nums[-1])} dopo 5 anni (base 100)"))
                break
    # V19: per le banche il divisore corretto e' il margine d'intermediazione
    if D.get("_bank"):
        ric, ni_alt = D.get("ricavi"), D.get("net_income")
        alt = D.get("ricavi_bancari")
        if alt and (ric is None or ric <= 1 or (ni_alt and ric < ni_alt)):
            D["ricavi"] = alt
            S["ricavi"] = [alt]
            log.append(("ok", f"ricavi        = {fmt(alt,0)} (margine di intermediazione)"))
    return D, S, log


# ============================================================================
# RILEVAMENTO NOME AZIENDA
# ============================================================================
BOILERPLATE = [r"united states", r"securities and exchange", r"commission",
               r"form 10-[kq]", r"annual report", r"transition report", r"proxy statement",
               r"commission file", r"exact name of registrant", r"address of registrant",
               r"zip code", r"telephone number", r"securities registered", r"indicate by check",
               r"accelerated filer", r"emerging growth", r"smaller reporting", r"shell company",
               r"act of 193", r"irs employer", r"identification no", r"for the fiscal year"]


def detect_company(lines):
    for i, ln in enumerate(lines[:80]):
        if "exact name of registrant" in ln.lower() and i > 0:
            for j in range(i - 1, -1, -1):
                cand = lines[j].strip()
                if cand and len(cand) > 2:
                    return cand[:90]
    freq = {}
    for ln in lines[:150]:
        for m in re.finditer(r"\b([A-Z][A-Za-z&'. -]{1,60}?(?:Corporation|Corp\.|Inc\.|S\.p\.A\.|S\.A\.|N\.V\.|plc|GmbH))\b", ln):
            name = m.group(1).strip()
            if len(name) > 3:
                freq[name] = freq.get(name, 0) + 1
    if freq:
        return max(freq.items(), key=lambda kv: kv[1])[0][:90]
    for ln in lines[:40]:
        low = ln.lower()
        if len(ln.strip()) < 4 or low.startswith("dear "):
            continue
        if any(re.search(p, low) for p in BOILERPLATE):
            continue
        return ln.strip()[:90]
    return lines[0][:90] if lines else "Report"


# ============================================================================
# UTILITÀ DI CALCOLO
# ============================================================================
def safe_div(a, b):
    if a is None or b is None or b == 0:
        return None
    return a / b


def cagr_of(series):
    if not series or len(series) < 2 or series[0] <= 0 or series[-1] <= 0:
        return None
    return (series[-1] / series[0]) ** (1.0 / (len(series) - 1)) - 1.0


def fmt(v, dec=1):
    if v is None:
        return "N/D"
    s = f"{v:,.{dec}f}"
    return s.replace(",", "§").replace(".", ",").replace("§", ".")


def fmt_pct(v, dec=1):
    return "N/D" if v is None else f"{fmt(v, dec)}%"


def status_from(value, pass_cond, warn_cond):
    if value is None:
        return "nd"
    if pass_cond:
        return "pass"
    if warn_cond:
        return "warn"
    return "fail"


# ============================================================================
# MOTORE DCF
# ============================================================================
class Metric:
    def __init__(self, code, name, sub, value, value_text, threshold, status, core, detail):
        self.code = code
        self.name = name
        self.sub = sub
        self.value = value
        self.value_text = value_text
        self.threshold = threshold
        self.status = status
        self.core = core
        self.detail = detail

    @property
    def tag(self):
        for dn in ("STATUS_META", "QSTATUS_META"):
            meta = globals().get(dn)
            if isinstance(meta, dict) and self.status in meta:
                return meta[self.status][1]
        return "… N/D"

    @property
    def color(self):
        for dn in ("STATUS_META", "QSTATUS_META"):
            meta = globals().get(dn)
            if isinstance(meta, dict) and self.status in meta:
                return meta[self.status][0]
        return "#8a93a6"


def compute_dcf(D, S):
    cfo, capex = D.get("cfo"), D.get("capex")
    oe = (cfo - capex) if (cfo is not None and capex is not None) else None
    base, base_is_oe = (oe, True) if (oe is not None and oe > 0) else (D.get("fcf"), False)
    if base is None or base <= 0:
        base, base_is_oe = D.get("net_income"), False
    if base is None or base <= 0:
        return {"ok": False}
    bv = D.get("equity")
    if bv is not None:
        bv = bv - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)
    use_bv = bool(D.get("_bank") and bv and bv > 0)
    roe_rep = D.get("roe_reported") or D.get("rote")
    if use_bv and roe_rep and D.get("net_income"):
        bv = D["net_income"] / roe_rep
    cagr_e = cagr_of(S.get("net_income"))
    g1_base = max(0.0, min(cagr_e, 0.08)) if cagr_e is not None else 0.04
    r, g, n = 0.10, 0.025, 10

    def val(base_v, g1, gp):
        if use_bv:
            roe_n = min((D.get("roe_reported") or D.get("rote") or safe_div(D.get("net_income"), bv) or 0), 0.25)
            flows, pv, bvt = [], 0.0, bv
            for t in range(1, n + 1):
                ri = (roe_n - r) * bvt
                flows.append(ri); pv += ri / (1 + r) ** t
                bvt *= (1 + gp)
            tv = (roe_n - r) * bvt / (r - gp)
            return bv + pv + tv / (1 + r) ** n, flows, tv
        flows, pv = [], 0.0
        for t in range(1, n + 1):
            f = base_v * (1 + g1) ** t
            flows.append(f); pv += f / (1 + r) ** t
        tv = flows[-1] * (1 + gp) / (r - gp)
        return pv + tv / (1 + r) ** n, flows, tv

    iv, flows, tv = val(base, g1_base, g)
    iv_bear, _, _ = val(base * 0.85, 0.0, 0.02)
    iv_bull, _, _ = val(base * 1.15, min(g1_base + 0.04, 0.12), 0.03)
    shares, price = D.get("shares"), D.get("price")
    mcap = D.get("market_cap")
    if price and shares:
        mc_calc = price * shares
        if mcap is None or not (0.5 * mc_calc <= mcap <= 2.0 * mc_calc):
            mcap = mc_calc
    ni_d = D.get("net_income")
    if mcap and ni_d and ni_d > 0 and not (2 <= mcap / ni_d <= 200):
        mcap = None
    iv_share = iv / shares if shares else None
    mos = ((iv - mcap) / iv * 100.0) if mcap else None
    mos_bear = ((iv_bear - mcap) / iv_bear * 100.0) if mcap else None
    mos_bull = ((iv_bull - mcap) / iv_bull * 100.0) if mcap else None
    ddm = (D.get("dividendi") * (1 + g) / (r - g)) if (use_bv and D.get("dividendi")) else None
    return {"ok": True, "base": base, "base_is_oe": base_is_oe, "base_is_bv": use_bv,
            "g1": g1_base, "r": r, "g": g, "n": n, "flows": flows, "tv": tv,
            "iv": iv, "iv_bear": iv_bear, "iv_bull": iv_bull,
            "mos": mos, "mos_bear": mos_bear, "mos_bull": mos_bull,
            "iv_share": iv_share, "mcap": mcap, "price": price, "cagr_used": cagr_e,
            "bv": bv if use_bv else None,
            "roe_bank": ((D.get("roe_reported") or D.get("rote") or safe_div(D.get("net_income"), bv)) if use_bv else None),
            "ddm": ddm}


def build_quant(D, S, dcf):
    Q = []
    def add(code, name, sub, val, vtxt, thr, st, core, det):
        Q.append(Metric(code, name, sub, val, vtxt, thr, st, core, det))
    oe, fcf = D.get("oe"), D.get("fcf")
    if oe is None: oe = fcf
    ni, eq = D.get("net_income"), D.get("equity")
    if D.get("_bank") and D.get("roe_reported") and ni:
        eq = ni / D["roe_reported"]
    add("Q01", "Owner Earnings", "Pilastro fondamentale · Cassa reale · CORE BUFFETT",
        oe, fmt(oe, 0), "Crescente e ≥ Utile Netto",
        status_from(oe, oe is not None and ni is not None and oe >= ni, oe is not None and oe > 0), True,
        f"CFO {fmt(D.get('cfo'),0)} − CapEx {fmt(D.get('capex'),0)}")
    nopat = (D.get("ebit") or 0) * (1 - (D.get("tax_rate") or 0.25))
    ic = (eq or 0) + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
    roic = safe_div(nopat, ic) if ic and ic > 0 else None
    add("Q02", "ROIC", "Efficienza del capitale · CORE BUFFETT",
        roic * 100 if roic else None, fmt_pct(roic * 100 if roic else None), "> 15% costante (Economic Moat)",
        status_from(roic * 100 if roic else None, roic is not None and roic > 0.15, roic is not None and roic > 0.10), True,
        f"NOPAT {fmt(nopat,0)} / Capitale investito {fmt(ic,0)}")
    roe = safe_div(ni, eq)
    roe_v = roe * 100 if roe is not None else None
    if D.get("_bank"):
        _rep = D.get("roe_reported") or D.get("rote")
        if _rep is not None and 0.05 <= _rep <= 0.60:
            roe = _rep; roe_v = _rep * 100
    oneoff = D.get("one_off")
    det3 = f"Utile {fmt(ni,0)} / Patrimonio Netto {fmt(eq,0)}"
    if oneoff and ni is not None and eq:
        det3 += f" · ROE rettificato one-off ({fmt(oneoff,0)}): {fmt((ni - oneoff) / eq * 100, 1)}%"
    _rp, _rw = (12, 10) if D.get("_bank") else (15, 10)
    add("Q03", "ROE", "Redditività dell'equity · CORE BUFFETT",
        roe_v, fmt_pct(roe_v), ("> 12% (soglia banche)" if D.get("_bank") else "> 15% costante (basso debito)"),
        status_from(roe_v, roe_v is not None and roe_v > _rp, roe_v is not None and roe_v >= _rw), True, det3)
    ltd, ni_ = D.get("lt_debt"), ni
    yrs = safe_div(ltd, ni_)
    add("Q04", "Anni per estinguere il debito", "Solvibilità · CORE BUFFETT",
        yrs, fmt(yrs, 2) + " anni" if yrs is not None else "N/D", "< 3–4 anni",
        status_from(yrs, yrs is not None and yrs < 3, yrs is not None and yrs < 4), True,
        f"Debito LT {fmt(ltd,0)} / Utile {fmt(ni_,0)}")
    de = safe_div(D.get("total_debt"), eq)
    add("Q05", "Debt / Equity", "Leva finanziaria · CORE BUFFETT",
        de, fmt(de, 2), "< 0,5",
        status_from(de, de is not None and de < 0.5, de is not None and de < 1.0), True,
        f"Debito {fmt(D.get('total_debt'),0)} / Equity {fmt(eq,0)}")
    gm = safe_div(D.get("gross_profit"), D.get("ricavi"))
    add("Q06", "Margine Lordo", "Pricing power · CORE BUFFETT",
        gm * 100 if gm else None, fmt_pct(gm * 100 if gm else None), "> 40% (vantaggio competitivo)",
        status_from(gm * 100 if gm else None, gm is not None and gm > 0.40, gm is not None and gm > 0.30), True,
        f"(Ricavi {fmt(D.get('ricavi'),0)} − COGS {fmt(D.get('cogs'),0)}) / Ricavi")
    nm = safe_div(ni, D.get("ricavi"))
    add("Q07", "Margine Netto", "Efficienza finale · CORE BUFFETT",
        nm * 100 if nm else None, fmt_pct(nm * 100 if nm else None), "> 15%–20%",
        status_from(nm * 100 if nm else None, nm is not None and nm > 0.15, nm is not None and nm > 0.10), True,
        f"Utile {fmt(ni,0)} / Ricavi {fmt(D.get('ricavi'),0)}")
    iv = dcf.get("iv")
    sh_q8 = D.get("shares")
    iv_ps = (iv / sh_q8) if (iv and sh_q8) else None
    add("Q08", "Valore Intrinseco (DCF Owner Earnings)", "Valutazione · CORE BUFFETT",
        (iv_ps if (D.get("_bank") and iv_ps) else iv), (fmt(iv_ps, 2) + " €/azione" if (D.get("_bank") and iv_ps) else fmt(iv, 0)), "Tasso sconto prudente 9–10%, g perpetua 2–3%",
        "pass" if iv else "nd", True,
        "" if iv is None else (f"Residual Income su patrimonio {fmt(dcf.get('bv'),0)} · ROE {fmt((dcf.get('roe_bank') or 0)*100,1)}% (normalizzato) · r=10% · g=2,5%" if dcf.get("base_is_bv") else f"Base {fmt(dcf.get('base'),0)} · r=10% · g=2,5% · n=10"))
    mcap, mos = dcf.get("mcap"), dcf.get("mos")
    add("Q09", "Margine di Sicurezza", "Protezione dal rischio · CORE BUFFETT",
        mos, fmt_pct(mos), "Sconto 20%–30% sul valore intrinseco",
        status_from(mos, mos is not None and mos >= 30, mos is not None and mos >= 20), True,
        f"(VI {fmt(iv,0)} − Market Cap {fmt(mcap,0)}) / VI" if iv and mcap else "")
    proxy = safe_div(mcap, eq)
    add("Q10", "Test del $1 di utile trattenuto", "Qualità del management · CORE BUFFETT",
        proxy, fmt(proxy, 2), "> $1,00 di valore per $1 trattenuto",
        status_from(proxy, proxy is not None and proxy > 1.5, proxy is not None and proxy > 1.0), True,
        f"Proxy Market Cap/Equity = {fmt(proxy,2)} (test completo richiede serie storica)")
    cr_capex = safe_div(D.get("capex"), fcf)
    add("Q11", "CapEx Ratio", "Business asset-light · CORE BUFFETT",
        cr_capex * 100 if cr_capex else None, fmt_pct(cr_capex * 100 if cr_capex else None), "< 25%–30%",
        status_from(cr_capex * 100 if cr_capex else None, cr_capex is not None and cr_capex < 0.25, cr_capex is not None and cr_capex < 0.30), True,
        f"CapEx {fmt(D.get('capex'),0)} / Utile {fmt(ni,0)}")
    add("Q12", "Share Buyback Rate", "Allocazione capitale · CORE BUFFETT", None, "N/D",
        "Riduzione azioni 1%–3% annuo", "nd", True, "Richiede serie storica azioni")
    sga_rl = safe_div(D.get("sga"), D.get("gross_profit"))
    add("Q13", "Spese SG&A / Utile Lordo", "Costi strutturali · CORE BUFFETT",
        sga_rl * 100 if sga_rl else None, fmt_pct(sga_rl * 100 if sga_rl else None), "< 30%",
        status_from(sga_rl * 100 if sga_rl else None, sga_rl is not None and sga_rl < 0.30, sga_rl is not None and sga_rl < 0.40), True,
        f"SGA {fmt(D.get('sga'),0)} / Margine Lordo {fmt(D.get('gross_profit'),0)}")
    rnd_rl = safe_div(D.get("rnd"), D.get("gross_profit"))
    add("Q14", "Spese R&D / Utile Lordo", "Rischio obsolescenza · CORE BUFFETT",
        rnd_rl * 100 if rnd_rl else None, fmt_pct(rnd_rl * 100 if rnd_rl else None), "< 15%",
        status_from(rnd_rl * 100 if rnd_rl else None, rnd_rl is not None and rnd_rl < 0.15, rnd_rl is not None and rnd_rl < 0.25), True,
        f"R&D {fmt(D.get('rnd'),0)} / Ricavi {fmt(D.get('ricavi'),0)}")
    ccr = safe_div(fcf, ni)
    add("Q15", "Cash Conversion Rate", "Qualità del bilancio · CORE BUFFETT",
        ccr, fmt(ccr, 2), "≥ 1,0 (assenza di artifizi contabili)",
        status_from(ccr, ccr is not None and ccr >= 1.0, ccr is not None and ccr >= 0.8), True,
        f"FCF {fmt(fcf,0)} / Utile {fmt(ni,0)}")
    eps = safe_div(ni, D.get("shares"))
    ey = safe_div(eps, D.get("price"))
    add("Q16", "Earnings Yield", "Rendimento azionario · CORE BUFFETT",
        ey * 100 if ey else None, fmt_pct(ey * 100 if ey else None), "Premio netto rispetto al risk-free",
        status_from(ey * 100 if ey else None, ey is not None and ey > 0.07, ey is not None and ey > 0.05), True,
        f"EPS {fmt(eps,2)} / Prezzo {fmt(D.get('price'),2)}")
    fcfy = safe_div(fcf, D.get("ev"))
    add("Q17", "Free Cash Flow Yield", "Rendimento di cassa · CORE BUFFETT",
        fcfy * 100 if fcfy else None, fmt_pct(fcfy * 100 if fcfy else None), "> 6%–7%",
        status_from(fcfy * 100 if fcfy else None, fcfy is not None and fcfy > 0.07, fcfy is not None and fcfy > 0.05), True,
        f"FCF {fmt(fcf,0)} / EV {fmt(D.get('ev'),0)}")
    roe_c = min(roe, 0.40) if roe is not None else None
    sgr = (roe_c * (1 - (D.get("payout") or 0))) if roe_c else None
    if sgr is not None: sgr = min(sgr, 0.25)
    add("Q18", "Sustainable Growth Rate", "Crescita interna · CORE BUFFETT",
        sgr * 100 if sgr else None, fmt_pct(sgr * 100 if sgr else None), "Elevato e coerente con lo storico",
        status_from(sgr * 100 if sgr else None, sgr is not None and sgr > 0.10, sgr is not None and sgr > 0.05), True,
        f"ROE {fmt(roe_v,1)}% × (1 − payout {fmt((D.get('payout') or 0)*100,1)}%)")
    add("Q19", "Retention Rate", "Reinvestimento utili · CORE BUFFETT",
        (D.get("payout") or 0) and 100 - (D.get("payout") or 0) * 100 or None,
        fmt_pct((100 - (D.get("payout") or 0) * 100) if D.get("payout") else None),
        "Dipende dalle opportunità ad alto ROIC",
        status_from((100 - (D.get("payout") or 0) * 100) if D.get("payout") else None, True, True), True,
        f"Payout {fmt((D.get('payout') or 0)*100,1)}%")
    fv = 10000 * (1 + (sgr or 0)) ** 10 if sgr else None
    add("Q20", "Crescita Composta (FV 10 anni)", "Effetto compounding · CORE BUFFETT",
        fv, fmt(fv, 0), "Focus sulla capitalizzazione di lungo periodo",
        status_from(fv, fv is not None and fv > 25000, fv is not None and fv > 20000), True,
        f"10.000 al tasso SGR {fmt((sgr or 0)*100,1)}% per 10 anni")
    add("Q21", "Buffett Indicator", "Macro valutazione di mercato", None, "N/D",
        "< 80% acquisto · > 120% sopravvalutato", "nd", False,
        "Dato macro (Market Cap/GDP) non presente nel report")
    om = safe_div(D.get("ebit"), D.get("ricavi"))
    add("Q22", "Margine Operativo", "Gestione caratteristica",
        om * 100 if om else None, fmt_pct(om * 100 if om else None), "> 20%",
        status_from(om * 100 if om else None, om is not None and om > 0.20, om is not None and om > 0.15), False,
        f"EBIT {fmt(D.get('ebit'),0)} / Ricavi {fmt(D.get('ricavi'),0)}")
    cg_n, cg_r = cagr_of(S.get("net_income")), cagr_of(S.get("ricavi"))
    add("Q23", "CAGR Utili / Ricavi", "Crescita storica composta",
        None, ("Ricavi " + fmt_pct(cg_r * 100 if cg_r else None) + " · Utili " + fmt_pct(cg_n * 100 if cg_n else None)) if (cg_r or cg_n) else "N/D",
        "> 5% annuo",
        status_from(max(cg_n or 0, cg_r or 0) * 100, (cg_n is not None and cg_n > 0.05) or (cg_r is not None and cg_r > 0.05), (cg_n is not None and cg_n > 0) or (cg_r is not None and cg_r > 0)), False,
        "Serie storica estratta dal report")
    add("Q24", "Free Cash Flow", "Liquidità", fcf, fmt(fcf, 0), "> 0",
        status_from(fcf, fcf is not None and fcf > 0, fcf is not None and fcf > 0), False,
        f"CFO {fmt(D.get('cfo'),0)} − CapEx {fmt(D.get('capex'),0)}")
    fm = safe_div(fcf, D.get("ricavi"))
    add("Q25", "FCF Margin", "Efficienza di cassa",
        fm * 100 if fm else None, fmt_pct(fm * 100 if fm else None), "> 10%",
        status_from(fm * 100 if fm else None, fm is not None and fm > 0.10, fm is not None and fm > 0.07), False, "")
    add("Q26", "EPS (Utile per Azione)", "Redditività azionaria",
        eps, fmt(eps, 2), "> 0 e crescente",
        status_from(eps, eps is not None and eps > 0, eps is not None and eps > 0), False,
        f"Utile {fmt(ni,0)} / Azioni {fmt(D.get('shares'),0)}")
    icr = safe_div(D.get("ebit"), D.get("interessi")) if D.get("interessi") else None
    add("Q27", "Interest Coverage", "Solvibilità interessi",
        icr, fmt(icr, 1) + "×" if icr else "N/D", "> 5 (ideale > 10)",
        status_from(icr, icr is not None and icr > 10, icr is not None and icr > 5), False,
        f"EBIT {fmt(D.get('ebit'),0)} / Interessi {fmt(D.get('interessi'),0)}" if D.get("interessi") else "Nessun debito oneroso rilevato")
    cur = safe_div(D.get("current_assets"), D.get("current_liab"))
    add("Q28", "Current Ratio", "Liquidità breve",
        cur, fmt(cur, 2), "> 1,5",
        status_from(cur, cur is not None and cur > 1.5, cur is not None and cur > 1.0), False,
        f"Attività correnti {fmt(D.get('current_assets'),0)} / Passività correnti {fmt(D.get('current_liab'),0)}")
    qr = safe_div((D.get("cassa") or 0) + (D.get("crediti") or 0), D.get("current_liab"))
    add("Q29", "Quick Ratio", "Liquidità immediata",
        qr, fmt(qr, 2), "≥ 1,0",
        status_from(qr, qr is not None and qr >= 1.0, qr is not None and qr >= 0.8), False,
        f"(Cassa {fmt(D.get('cassa'),0)} + Crediti {fmt(D.get('crediti'),0)}) / Pass. correnti {fmt(D.get('current_liab'),0)}")
    roa = safe_div(ni, D.get("total_assets"))
    _rp30, _rw30 = (1.0, 0.7) if D.get("_bank") else (10, 5)
    add("Q30", "ROA", "Efficienza sugli asset",
        roa * 100 if roa else None, fmt_pct(roa * 100 if roa else None),
        ("> 1,0% (soglia banche)" if D.get("_bank") else "> 10%"),
        status_from(roa * 100 if roa else None, roa is not None and roa * 100 > _rp30, roa is not None and roa * 100 > _rw30), False, "")
    roce = safe_div(D.get("ebit"), (D.get("total_assets") or 0) - (D.get("current_liab") or 0))
    add("Q31", "ROCE", "Capitale impiegato",
        roce * 100 if roce else None, fmt_pct(roce * 100 if roce else None), "> 15%",
        status_from(roce * 100 if roce else None, roce is not None and roce > 0.15, roce is not None and roce > 0.10), False,
        f"EBIT / (Totale attivo {fmt(D.get('total_assets'),0)} − Pass. corr. {fmt(D.get('current_liab'),0)})")
    pe = safe_div(D.get("price"), eps)
    add("Q32", "P/E", "Multiplo di mercato",
        pe, fmt(pe, 1), "< 15 interessante",
        status_from(pe, pe is not None and pe < 15, pe is not None and pe < 20), False,
        f"Prezzo {fmt(D.get('price'),2)} / EPS {fmt(eps,2)}")
    peg = safe_div(pe, (cg_n or 0) * 100) if pe else None
    add("Q33", "PEG Ratio", "Multiplo corretto per crescita",
        peg, fmt(peg, 2), "< 1,0 sottovalutazione relativa",
        status_from(peg, peg is not None and peg < 1.0, peg is not None and peg < 1.5), False,
        f"P/E {fmt(pe,1)} / crescita utili {fmt_pct(cg_n*100 if cg_n else None)}")
    if D.get("_bank") and D.get("roe_reported") and ni:
        eq = ni / D["roe_reported"]
    bvps = safe_div(eq, D.get("shares"))
    pb = safe_div(D.get("price"), bvps)
    add("Q34", "Price / Book Value", "Multiplo contabile",
        pb, fmt(pb, 2), "< 3 (contesto ROE elevato)",
        status_from(pb, pb is not None and pb < 3, pb is not None and pb < 5), False,
        f"Prezzo {fmt(D.get('price'),2)} / BVPS {fmt(bvps,2)}")
    evebit = safe_div(D.get("ev"), D.get("ebit"))
    add("Q35", "EV / EBIT", "Valutazione globale",
        evebit, fmt(evebit, 1), "< 12",
        status_from(evebit, evebit is not None and evebit < 12, evebit is not None and evebit < 15), False,
        f"EV {fmt(D.get('ev'),0)} / EBIT {fmt(D.get('ebit'),0)}")
    evfcf = safe_div(D.get("ev"), fcf)
    add("Q36", "EV / FCF", "Valutazione su cassa",
        evfcf, fmt(evfcf, 1), "< 15",
        status_from(evfcf, evfcf is not None and evfcf < 15, evfcf is not None and evfcf < 20), False,
        f"EV {fmt(D.get('ev'),0)} / FCF {fmt(fcf,0)}")
    add("Q37", "Dividend Payout Ratio", "Politica dividendi",
        (D.get("payout") or 0) * 100 if D.get("payout") else None,
        fmt_pct((D.get("payout") or 0) * 100 if D.get("payout") else None), "< 60% sostenibile",
        status_from((D.get("payout") or 0) * 100 if D.get("payout") else None, (D.get("payout") or 0) < 0.6, (D.get("payout") or 0) < 0.8), False, "")
    dr = safe_div(D.get("total_debt"), D.get("total_assets"))
    add("Q38", "Debt Ratio", "Struttura patrimoniale",
        dr * 100 if dr else None, fmt_pct(dr * 100 if dr else None), "< 40%",
        status_from(dr * 100 if dr else None, dr is not None and dr < 0.40, dr is not None and dr < 0.60), False, "")
    dde = safe_div(D.get("lt_debt"), eq)
    add("Q39", "Debito Oneroso / Equity", "Peso del debito finanziario",
        dde, fmt(dde, 2), "< 0,5",
        status_from(dde, dde is not None and dde < 0.5, dde is not None and dde < 1.0), False, "")
    add("Q40", "CAGR Book Value per Azione", "Storico patrimoniale", None, "N/D",
        "Crescita composta del BVPS", "nd", False, "Richiede serie storica pluriennale")
    if D.get("_bank"):
        for m in Q:
            if m.code in ("Q01", "Q02", "Q04", "Q05", "Q06", "Q07", "Q11", "Q13", "Q14", "Q15", "Q17", "Q22", "Q24", "Q25", "Q27", "Q28", "Q29", "Q31", "Q35", "Q36", "Q38", "Q39"):
                m.status = "nd"; m.hidden = True
        tbv = (eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)) if eq is not None else None
        sh, pr = D.get("shares"), D.get("price")
        ptbv = (pr / (tbv / sh)) if (pr and sh and tbv and tbv > 0) else None
        cet1, ci, npl = D.get("cet1"), D.get("cost_income"), D.get("npl")
        if roe_v is not None:
            add("B1", "ROE banche", "Metrica regina di Buffett per le banche", roe_v, fmt_pct(roe_v),
                "> 12–15%", status_from(roe_v, roe_v > 12, roe_v > 10), True, f"Utile {fmt(ni,0)} / Equity {fmt(eq,0)}")
        if roa is not None:
            add("B2", "ROA banche", "Return on assets (soglia bancaria)", roa*100, fmt_pct(roa*100),
                "> 1,0%", status_from(roa*100, roa*100 > 1.0, roa*100 > 0.7), True, f"Utile {fmt(ni,0)} / Attivo {fmt(D.get('total_assets'),0)}")
        if ptbv is not None:
            add("B3", "P/TBV", "Prezzo / tangible book value", ptbv, fmt(ptbv,2),
                "< 1,5 (sconto su tangibile)", status_from(ptbv, ptbv < 1.5, ptbv < 2.0), True,
                f"Prezzo {fmt(pr,2)} / TBV/azione {fmt(tbv/sh,2)}")
        if cet1 is not None:
            add("B4", "CET1 Ratio", "Solidità patrimoniale", cet1*100, fmt_pct(cet1*100),
                "> 13%", status_from(cet1*100, cet1*100 > 13, cet1*100 > 11.5), True, "reportistica prudenziale")
        if ci is not None:
            add("B5", "Cost / Income", "Efficienza operativa", ci*100, fmt_pct(ci*100),
                "< 50–55%", status_from(ci*100, ci*100 < 55, ci*100 < 65), True, "costi operativi / margine d'intermediazione")
        if True:
            npl_v, loans_v = D.get("npl"), D.get("loans")
            npl_pct = None
            if npl_v is not None and npl_v <= 1.0:
                npl_pct = npl_v * 100
            elif npl_v is not None and loans_v and loans_v > 0:
                npl_pct = npl_v / loans_v * 100
            if npl_pct is not None:
                add("B6", "NPL Ratio", "Qualità dei crediti", npl_pct, fmt_pct(npl_pct),
                    "< 3%", status_from(npl_pct, npl_pct < 3, npl_pct < 5), True,
                    (f"Sofferenze lorde {fmt(npl_v,0)} / Crediti verso clientela {fmt(loans_v,0)}" if loans_v else "NPL ratio da reportistica"))
    if D.get("_bank"):
        for m in Q:
            if m.code in ("Q01","Q02","Q03","Q04","Q05","Q06","Q07","Q11","Q13","Q14","Q15","Q17",
                          "Q22","Q24","Q25","Q27","Q28","Q29","Q30","Q31","Q35","Q36","Q38","Q39"):
                m.status = "nd"; m.hidden = True
            if "Non significativo" in (m.detail or ""):
                m.hidden = True
# Q23 bank nd
    if D.get("_bank"):
        sn = S.get("net_income") or []
        for m in Q:
            if m.code == "Q23" and len(sn) < 2:
                m.status = "nd"; m.hidden = True
    return Q


def build_qual(D, S, dcf):
    K = []
    ric, cogs, ni, ebit = D.get("ricavi"), D.get("cogs"), D.get("net_income"), D.get("ebit")
    gp = D.get("gross_profit")
    eq, ltd, ta = D.get("equity"), D.get("lt_debt"), D.get("total_assets")
    cfo, capex, fcf = D.get("cfo"), D.get("capex"), D.get("fcf")
    sga, rnd = D.get("sga"), D.get("rnd")
    mcap = D.get("market_cap") or ((D.get("price") or 0) * (D.get("shares") or 0) or None)
    tax = D.get("tax_rate", 0.24)

    gm = safe_div(gp, ric)
    gm_v = gm * 100 if gm is not None else None
    roic = safe_div(ebit * (1 - tax) if ebit is not None else None,
                    (eq + (D.get("total_debt") or 0) - (D.get("cassa") or 0)) if eq is not None else None)
    roic_v = roic * 100 if roic is not None else None

    def add(n, title, desc, formula, status, note):
        K.append({"n": n, "title": title, "desc": desc, "formula": formula,
                  "status": status, "note": note})

    if gm_v is None and roic_v is None:
        st = "nd"; note = "Dati insufficienti"
    else:
        ok = (gm_v is not None and gm_v > 40) and (roic_v is not None and roic_v > 15)
        part = (gm_v is not None and gm_v > 40) or (roic_v is not None and roic_v > 15)
        st = "pass" if ok else ("partial" if part else "fail")
        note = f"Margine Lordo {fmt_pct(gm_v)} · ROIC {fmt_pct(roic_v)}"
    add(1, "Vantaggio competitivo durevole (Economic Moat)",
        "Fossato protettivo (brand, network effect, switching cost, vantaggio di costo) che impedisce ai concorrenti di erodere i profitti.",
        "Margine Lordo > 40% · ROIC > 15%", st, note)

    sr, se = S.get("ricavi"), S.get("ebit")
    if sr and se and len(sr) >= 2 and len(se) >= 2 and len(sr) == len(se):
        margins = [e / r * 100 for r, e in zip(sr, se) if r]
        mean = sum(margins) / len(margins)
        if len(margins) > 1:
            std = math.sqrt(sum((x - mean) ** 2 for x in margins) / (len(margins) - 1))
        else:
            std = 0.0
        st = "pass" if (std < 2 and mean > 20) else ("partial" if mean > 15 else "fail")
        note = f"Margini operativi {len(margins)} anni: media {fmt(mean,1)}%, deviazione std {fmt(std,2)} pp"
    else:
        om = safe_div(ebit, ric)
        om_v = om * 100 if om is not None else None
        st = "partial" if (om_v is not None and om_v > 20) else ("fail" if om_v is not None else "nd")
        note = f"Margine operativo {fmt_pct(om_v)} (stabilità pluriennale non verificabile senza serie storica)"
    add(2, "Potere di prezzo (Pricing Power)",
        "Capacità di alzare i prezzi senza perdere clienti a favore della concorrenza.",
        "Margine Operativo = EBIT / Ricavi, stabile nel tempo (σ ≈ 0)", st, note)

    sni = S.get("net_income")
    if sni and len(sni) >= 2:
        neg = sum(1 for x in sni if x < 0)
        st = "pass" if neg == 0 else "fail"
        note = f"Serie utili {len(sni)} anni: {neg} esercizi in perdita"
    elif ni is not None:
        st = "partial"; note = f"Utile corrente positivo ({fmt(ni,0)}), ma serve serie storica per la continuità"
    else:
        st = "nd"; note = "Dato non disponibile"
    add(3, "Cerchio di competenza e prevedibilità",
        "Business semplice, domanda costante, futuro prevedibile a 10–20 anni.",
        "Anni in perdita negli ultimi 10 = 0", st, note)

    cx = safe_div(capex, cfo)
    cx_v = cx * 100 if cx is not None else None
    add(4, "Modello capital-light",
        "Il business cresce senza reinvestire continuamente ingenti capitali in impianti e magazzino.",
        "CapEx / Cash Flow Operativo < 25%–30%",
        status_from(cx_v, cx_v is not None and cx_v <= 30, cx_v is not None and cx_v <= 50),
        f"CapEx {fmt(capex,0)} / CFO {fmt(cfo,0)} = {fmt_pct(cx_v)}")

    oe = (cfo - capex) if (cfo is not None and capex is not None) else None
    cc = safe_div(fcf, ni)
    if oe is not None and cc is not None:
        st = "pass" if (oe > 0 and cc >= 1.0) else ("partial" if oe > 0 else "fail")
    else:
        st = "nd"
    add(5, "Generazione reale di cassa (Owner Earnings)",
        "Gli utili contabili devono trasformarsi in denaro reale e liquido per gli azionisti.",
        "Owner Earnings = CFO − CapEx · Cash Conversion ≥ 100%", st,
        f"Owner Earnings {fmt(oe,0)} · Conversione cassa {fmt(cc,2)}")

    proxy = safe_div(mcap, eq)
    if proxy is not None:
        st = "pass" if proxy >= 1.5 else "partial"
        note = f"Proxy Market Cap/Equity = {fmt(proxy,2)} (il test completo richiede la serie storica a 10 anni)"
    else:
        st = "nd"; note = "Richiede serie storica di market cap e utili trattenuti"
    add(6, "Allocazione razionale del capitale",
        "Management che reinveste ad alto rendimento, buyback a sconto, dividendi sostenibili.",
        "Test del $1: Δ Market Cap / Utili trattenuti > 1,0", st, note)

    den = gp if gp else ric
    sg = safe_div(sga, den)
    sg_v = sg * 100 if sg is not None else None
    add(7, "Frugalità e controllo dei costi",
        "Cultura aziendale orientata all'efficienza e al contenimento delle spese generali.",
        "SG&A / Utile Lordo < 30%",
        status_from(sg_v, sg_v is not None and sg_v < 30, sg_v is not None and sg_v <= 40),
        f"SGA {fmt(sga,0)} / {'Margine Lordo' if gp else 'Ricavi'} {fmt(den,0)} = {fmt_pct(sg_v)}")

    roa = safe_div(ni, ta)
    roa_v = roa * 100 if roa is not None else None
    add(8, "Criterio dell'idiot test",
        "Business così forte da generare profitti anche se gestito da un incompetente.",
        "ROA = Utile Netto / Totale Attività > 10%",
        status_from(roa_v, roa_v is not None and roa_v > 10, roa_v is not None and roa_v >= 5),
        f"ROA = {fmt_pct(roa_v)}")

    yrs = safe_div(ltd, ni)
    add(9, "Indipendenza e avversione al debito",
        "Capacità di superare recessioni e tassi alti senza debito rischioso.",
        "Debito LT / Utile Netto < 3–4 anni",
        status_from(yrs, yrs is not None and yrs < 3, yrs is not None and yrs <= 4),
        f"{fmt(yrs,2)} anni di utili per estinguere il debito LT")

    nm = safe_div(ni, ric)
    nm_v = nm * 100 if nm is not None else None
    add(10, "Modello franchise vs commodity",
        "Prodotto unico e insostituibile, non merce indifferenziata da guerra di prezzo.",
        "Margine Netto > 15%",
        status_from(nm_v, nm_v is not None and nm_v > 15, nm_v is not None and nm_v >= 10),
        f"Margine netto = {fmt_pct(nm_v)}")

    rr = safe_div(rnd, ric)
    rr_v = rr * 100 if rr is not None else None
    if rnd is None and ric:
        st, note = "warn", "⚠ R&D non esplicitata nel report (dato mancante, non assume valore 0)"
    else:
        st = status_from(rr_v, rr_v is not None and rr_v <= 3, rr_v is not None and rr_v <= 8)
        note = f"R&D {fmt(rnd,0)} / Ricavi {fmt(ric,0)} = {fmt_pct(rr_v)}"
    add(11, "Resistenza a disruption e obsolescenza",
        "Bisogno umano primario, basso rischio di obsolescenza tecnologica.",
        "R&D / Ricavi molto bassa o assente", st, note)

    ratio = safe_div(cfo, ebit)
    ta_q = D.get("total_assets")
    accrual = safe_div((ni or 0) - (cfo or 0), ta_q) if (ni is not None and cfo is not None and ta_q) else None
    checks, notes = [], []
    if ratio is not None:
        checks.append((0.8 <= ratio <= 1.3) or (bool(D.get("_bank")) and 0.5 <= ratio <= 1.5))
        notes.append(f"CFO/EBIT {fmt(ratio,2)}")
    if accrual is not None:
        checks.append(abs(accrual) < 0.10)
        notes.append(f"accrual ratio {fmt(accrual * 100,1)}%")
    checks.append(not D.get("_restatement"))
    notes.append("nessun restatement" if not D.get("_restatement") else "RESTATEMENT rilevato")
    if D.get("_unqualified") is not None:
        checks.append(bool(D.get("_unqualified")))
        notes.append("auditor unqualified" if D.get("_unqualified") else "auditor NON unqualified")
    sc12 = sum(checks) / len(checks) if checks else 0
    st12 = "pass" if sc12 >= 0.75 else ("partial" if sc12 >= 0.5 else "fail")
    add(12, "Candore, trasparenza e onestà contabile",
        "Management senza ego, niente artifizi contabili o EBITDA adjusted fantasiosi.",
        "CFO/EBIT ≈ 1,0 · accrual < 10% · nessun restatement · auditor unqualified",
        st12, " · ".join(notes))
    if D.get("_bank"):
        for q in K:
            if q["n"] in (4, 5, 9, 12):
                q["status"] = "nd"
                q["note"] = q["note"] + " · Non significativo per istituti finanziari"
    return K


# ============================================================================
# PUNTEGGI E VERDETTO
# ============================================================================
QUANT_AREAS = [
    ("Redditivita & efficienza", 0.35, ["Q02","Q03","Q06","Q07","Q13","Q14","Q22","Q26","Q30","Q31"]),
    ("Solidita & leva", 0.25, ["Q04","Q05","Q27","Q28","Q29","Q38","Q39"]),
    ("Cassa & crescita", 0.25, ["Q01","Q11","Q12","Q15","Q18","Q19","Q20","Q23","Q24","Q25"]),
    ("Valutazione & mercato", 0.15, ["Q08","Q09","Q10","Q16","Q17","Q21","Q32","Q33","Q34","Q35","Q36","Q37","Q40"]),
]
QUAL_AREAS = [
    ("Moat & durata", 0.40, [1, 2, 10, 11]),
    ("Management & capitale", 0.30, [6, 7, 12]),
    ("Prevedibilita & modello", 0.30, [3, 4, 5, 8, 9]),
]


def descriptor_for(f):
    if f >= 90: return "Qualita eccezionale: moat profondo, numeri da compounding machine"
    if f >= 80: return "Qualita alta: business solido con vantaggi competitivi chiari"
    if f >= 70: return "Qualita buona: fondamenta solide, alcune aree da monitorare"
    if f >= 60: return "Qualita discreta: criticita significative presenti"
    if f >= 50: return "Qualita mediocre: non supera pienamente i filtri di Buffett"
    return "Qualita insufficiente: business non compatibile con il metodo"


def compute_adv_proxies(text, D, S):
    """v2.2: proxy qualitativi avanzati, estratti solo dal testo (spie di 2° livello)."""
    adv = []
    L = len(text) or 1

    ng = len(re.findall(r"\bnon[- ]GAAP\b|\badjusted\b|una tantum|non ricorrent|one[- ]time charge|excluding\b|al netto di\b", text, re.I))
    dens = ng / (L / 100000.0)
    if ng == 0:
        st, nt = "pass", "Nessun linguaggio adjusted/non-GAAP rilevato: contabilità pulita"
    elif dens < 15:
        st, nt = "pass", f"Linguaggio adjusted/non-GAAP limitato ({ng} occorrenze, densità {dens:.0f}/100k caratteri)"
    elif dens < 30:
        st, nt = "warn", f"Uso moderato di adjusted/non-GAAP ({ng} occorrenze, densità {dens:.0f}/100k): verificare cosa viene escluso"
    else:
        st, nt = "fail", f"Forte ricorso ad adjusted/non-GAAP ({ng} occorrenze, densità {dens:.0f}/100k): leggere cosa viene escluso"
    adv.append({"name": "A1 · Linguaggio adjusted / non-GAAP", "status": st, "note": nt})

    pos = len(re.findall(r"(?:exceeded|surpassed|ahead of|above) (?:our|its|the) (?:guidance|expectations|estimates)", text, re.I))
    neg = len(re.findall(r"(?:below|missed|short of|behind) (?:our|its|the) (?:guidance|expectations|estimates)", text, re.I))
    if pos == 0 and neg == 0:
        st, nt = "nd", "Nessuna dichiarazione esplicita guidance-vs-actual trovata"
    elif neg == 0:
        st, nt = "pass", f"Solo dichiarazioni positive sulla guidance ({pos} superamenti, 0 miss)"
    elif neg <= pos:
        st, nt = "warn", f"Bilancio misto: {pos} superamenti vs {neg} miss dichiarati"
    else:
        st, nt = "fail", f"Prevalenza di miss vs guidance ({neg} negativi vs {pos} positivi)"
    adv.append({"name": "A2 · Coerenza guidance vs risultati", "status": st, "note": nt})

    rpo, ric = D.get("rpo"), D.get("ricavi")
    if rpo and ric and ric > 0:
        ratio = rpo / ric
        st = "pass" if ratio >= 0.75 else ("warn" if ratio >= 0.35 else "fail")
        nt = f"RPO/backlog = {fmt(ratio,2)}x i ricavi annui: visibilità {'alta' if ratio >= 0.75 else 'media' if ratio >= 0.35 else 'bassa'} sui ricavi futuri"
    else:
        st, nt = "nd", "RPO/backlog non presenti (tipico di chi vende a transazione, non in abbonamento)"
    adv.append({"name": "A3 · Visibilità ricavi (RPO/backlog)", "status": st, "note": nt})

    m = re.search(r"dear shareholders", text, re.I)
    if m:
        letter = text[m.start(): m.start() + 8000]
        negw = len(re.findall(r"(?:challeng|difficult|declin|decreas|adverse|setback|headwind|lower than|uncertain|risk)\w*", letter, re.I))
        posw = len(re.findall(r"(?:record|strong|grew|increas|surpass|accelerat|momentum|confident)\w*", letter, re.I))
        if negw >= 3:
            st, nt = "pass", f"Lettera bilanciata: {negw} ammissioni di difficoltà vs {posw} messaggi positivi (candore)"
        elif negw >= 1:
            st, nt = "warn", f"Lettera poco autocritica: {negw} menzioni di difficoltà vs {posw} positivi"
        else:
            st, nt = "fail", f"Lettera solo autocelebrativa: 0 difficoltà ammesse vs {posw} positivi"
    else:
        st, nt = "nd", "Lettera agli azionisti non trovata nel documento"
    adv.append({"name": "A4 · Tono/candore lettera agli azionisti", "status": st, "note": nt})

    ni, cfo, ta = D.get("net_income"), D.get("cfo"), D.get("total_assets")
    if ni is not None and cfo is not None and ta:
        accr = (ni - cfo) / ta
        st = "pass" if abs(accr) < 0.10 else ("warn" if abs(accr) < 0.20 else "fail")
        nt = f"Accrual ratio {fmt(accr * 100, 1)}% (utile vs cassa su attivo): " + ("utile di qualità, supportato dalla cassa" if abs(accr) < 0.10 else "utile poco supportato dalla cassa: approfondire")
        adv.append({"name": "A5 · Accrual ratio (qualità dell'utile)", "status": st, "note": nt})
    return adv


def compute_scores(quant, qual):
    def pts(st):
        if st == "pass": return 1.0
        if st in ("warn", "partial"): return 0.5
        if st == "fail": return 0.0
        return None
    qmap = {m.code: m for m in quant}
    kmap = {q["n"]: q for q in qual}

    def area_score(codes, getter):
        vals = [v for v in (getter(c) for c in codes) if v is not None]
        return (100.0 * sum(vals) / len(vals)) if vals else None

    q_areas, acc, wacc = [], 0.0, 0.0
    for name, w, codes in QUANT_AREAS:
        v = area_score(codes, lambda c: pts(qmap[c].status) if c in qmap else None)
        if v is not None:
            q_areas.append((name, w, v)); acc += w * v; wacc += w
    quant_score = acc / wacc if wacc else 0.0

    k_areas, acc, wacc = [], 0.0, 0.0
    for name, w, codes in QUAL_AREAS:
        v = area_score(codes, lambda c: pts(kmap[c]["status"]) if c in kmap else None)
        if v is not None:
            k_areas.append((name, w, v)); acc += w * v; wacc += w
    qual_score = acc / wacc if wacc else 0.0

    final = 0.60 * qual_score + 0.40 * quant_score
    qv = [m for m in quant if m.status != "nd"]
    ka = [q for q in qual if q["status"] != "nd"]
    return {"quant": quant_score, "qual": qual_score, "final": final,
            "quant_n": len(qv), "qual_n": len(ka),
            "qual_pass": sum(1 for q in ka if q["status"] == "pass"),
            "q_areas": q_areas, "k_areas": k_areas}


def verdict_for(scores, dcf):
    f = scores["final"]
    if f >= 85:
        t = "🏆 ECCELLENTE — Qualità da cassaforte Berkshire"
    elif f >= 70:
        t = "✅ BUONO — Azienda solida, da approfondire"
    elif f >= 55:
        t = "⚠️ MEDIOCRE — Non supera pienamente i filtri di Buffett"
    else:
        t = "❌ SCARTARE — Non compatibile con il metodo"
    mos = dcf.get("mos") if dcf.get("ok") else None
    if mos is None:
        sub = "Prezzo di mercato non disponibile nel filing: il margine di sicurezza va verificato con il prezzo corrente di borsa."
    elif mos > 95:
        sub = f"Margine di sicurezza {fmt(mos,1)}%: ANOMALO (>95%). Verificare coerenza prezzo/azioni nel LOG."
    elif mos >= 30:
        sub = f"Margine di sicurezza {fmt(mos,1)}%: prezzo interessante, sconto adeguato (≥ 30%)."
    elif mos >= 20:
        sub = f"Margine di sicurezza {fmt(mos,1)}%: sconto borderline (20–30%), serve prudenza."
    elif mos >= 0:
        sub = f"Margine di sicurezza {fmt(mos,1)}%: prezzo vicino al valore intrinseco. Buffett attenderebbe uno sconto maggiore."
    else:
        sub = f"Prezzo SOPRA il valore intrinseco ({fmt(mos,1)}%): attendere una correzione."
    return t, sub


# ============================================================================
# PIPELINE DI ANALISI
# ============================================================================
def analyze_document(path, cb=None):
    def step(p, msg):
        if cb:
            cb(p, msg)

    step(5, "Lettura del documento…")
    text = extract_text_from_file(path)
    try:
        with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), "w", encoding="utf-8") as f:
            f.write(text)
    except Exception:
        pass
    if not text.strip():
        raise ValueError("Impossibile estrarre testo dal documento.")

    step(25, "Estrazione automatica dei dati finanziari…")
    D, S, log = extract_financials(text)
    found = sum(1 for k in D if k != "_unit")
    if found < 3:
        raise ValueError("Nessun dato finanziario significativo riconosciuto nel documento.")

    step(55, "Calcolo dei 40 criteri quantitativi…")
    dcf = compute_dcf(D, S)
    quant = build_quant(D, S, dcf)

    step(75, "Analisi dei 12 criteri qualitativi…")
    qual = build_qual(D, S, dcf)

    step(90, "Punteggi e verdetto finale…")
    scores = compute_scores(quant, qual)
    vt, vs = verdict_for(scores, dcf)

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    company = detect_company(lines)
    if re.search(r"\b(bank|banca|insurance|assicur|credito)\b", company, re.I):
        D["_bank"] = 1
    step(100, "Completato.")
    if "adv" not in locals():
        try:
            adv = compute_adv_proxies(text, D, S)
        except Exception:
            adv = []
    return {"source": os.path.basename(path), "company": company,
            "text_len": len(text), "D": D, "S": S, "quant": quant, "qual": qual,
            "dcf": dcf, "scores": scores, "verdict": vt, "verdict_sub": vs, "adv": adv, "log": log}


# ============================================================================
# GUI — WIDGET HELPER
# ============================================================================
class ScrollableFrame(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent, bg=BG)
        self.canvas = tk.Canvas(self, bg=BG, highlightthickness=0, bd=0)
        self.scroll = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.inner = tk.Frame(self.canvas, bg=BG)
        self.inner.bind("<Configure>",
                        lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.win = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scroll.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scroll.pack(side="right", fill="y")
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfigure(self.win, width=max(1, e.width - 20)))
        self.canvas.bind("<Enter>", self._bind_wheel)
        self.canvas.bind("<Leave>", self._unbind_wheel)

    def _wheel(self, event):
        if event.num == 4:
            self.canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            self.canvas.yview_scroll(1, "units")
        else:
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _bind_wheel(self, _):
        self.canvas.bind_all("<MouseWheel>", self._wheel)
        self.canvas.bind_all("<Button-4>", self._wheel)
        self.canvas.bind_all("<Button-5>", self._wheel)

    def _unbind_wheel(self, _):
        self.canvas.unbind_all("<MouseWheel>")
        self.canvas.unbind_all("<Button-4>")
        self.canvas.unbind_all("<Button-5>")


def clear(frame):
    for w in frame.winfo_children():
        w.destroy()


def pill(parent, status, meta=None):
    meta = meta or STATUS_META
    col, lab = meta.get(status, (GRAY, "… N/D"))
    return tk.Label(parent, text=lab, bg=col, fg="#0a0f16",
                    font=("Segoe UI", 8, "bold"), padx=8, pady=2)


def metric_card(parent, m):
    card = tk.Frame(parent, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
    col = STATUS_META.get(m.status, (GRAY, "… N/D"))[0]
    bar = tk.Frame(card, bg=col, width=4)
    bar.grid(row=0, column=0, rowspan=5, sticky="ns")
    head = tk.Frame(card, bg=CARD)
    head.grid(row=0, column=1, sticky="ew", padx=(12, 10), pady=(10, 0))
    tk.Label(head, text=f"{m.code} · {m.name}", bg=CARD, fg=FG, font=F_CARD_T, anchor="w").pack(side="left")
    tk.Label(head, text="CORE BUFFETT" if m.core else "COMPLEMENTARE", bg=CARD,
             fg=GOLD if m.core else MUTED, font=F_TINY).pack(side="right")
    tk.Label(card, text=m.tag, bg=CARD, fg=MUTED, font=F_TINY, anchor="w").grid(
        row=1, column=1, sticky="ew", padx=(12, 10))
    vrow = tk.Frame(card, bg=CARD)
    vrow.grid(row=2, column=1, sticky="ew", padx=(12, 10), pady=(4, 0))
    tk.Label(vrow, text=m.value_text, bg=CARD, fg=FG, font=("Segoe UI", 15, "bold"), anchor="w").pack(side="left")
    pill(vrow, m.status).pack(side="right")
    if m.detail:
        tk.Label(card, text=m.detail, bg=CARD, fg=BLUE, font=F_TINY, anchor="w").grid(
            row=3, column=1, sticky="ew", padx=(12, 10), pady=(2, 0))
    tk.Label(card, text=f"Soglia Buffett: {m.threshold}", bg=CARD, fg=MUTED, font=F_TINY, anchor="w").grid(
        row=4, column=1, sticky="ew", padx=(12, 10), pady=(2, 10))
    return card


def qual_card(parent, q):
    card = tk.Frame(parent, bg=CARD, highlightbackground=BORDER, highlightthickness=1, padx=16, pady=12)
    col, lab = QSTATUS_META.get(q["status"], (GRAY, "… N/D"))
    top = tk.Frame(card, bg=CARD)
    top.pack(fill="x")
    tk.Label(top, text=f"{q['n']:02d} · {q['title']}", bg=CARD, fg=FG, font=("Segoe UI", 11, "bold"),
             anchor="w").pack(side="left")
    pill(top, q["status"], QSTATUS_META).pack(side="right")
    tk.Label(card, text=q["desc"], bg=CARD, fg=MUTED, font=F_SMALL, anchor="w",
             justify="left", wraplength=620).pack(fill="x", pady=(6, 4))
    tk.Label(card, text="📐 " + q["formula"], bg=CARD, fg=GOLD, font=F_SMALL, anchor="w").pack(fill="x")
    tk.Label(card, text="→ " + q["note"], bg=CARD, fg=col, font=F_SMALL, anchor="w",
             justify="left", wraplength=620).pack(fill="x", pady=(4, 0))
    return card


def kpi_card(parent, label, value, sub="", color=FG):
    c = tk.Frame(parent, bg=CARD2, highlightbackground=BORDER, highlightthickness=1, padx=14, pady=10)
    tk.Label(c, text=label, bg=CARD2, fg=MUTED, font=F_TINY, anchor="w").pack(fill="x")
    tk.Label(c, text=value, bg=CARD2, fg=color, font=("Segoe UI", 16, "bold"), anchor="w").pack(fill="x")
    if sub:
        tk.Label(c, text=sub, bg=CARD2, fg=MUTED, font=F_TINY, anchor="w").pack(fill="x")
    return c


def draw_gauge(cv, score, size=230):
    cv.delete("all")
    cx = cy = size // 2
    r = size // 2 - 20
    cv.create_arc(cx - r, cy - r, cx + r, cy + r, start=240, extent=-240,
                  style="arc", outline=BORDER, width=16)
    color = GREEN if score >= 80 else (GOLD if score >= 65 else (AMBER if score >= 50 else RED))
    ext = -240 * max(0.0, min(score, 100.0)) / 100.0
    if ext < -1:
        cv.create_arc(cx - r, cy - r, cx + r, cy + r, start=240, extent=ext,
                      style="arc", outline=color, width=16)
    cv.create_text(cx, cy - 12, text=f"{score:.0f}", font=("Segoe UI", 40, "bold"), fill=FG)
    cv.create_text(cx, cy + 26, text="/ 100", font=("Segoe UI", 11), fill=MUTED)
    cv.create_text(cx, cy + 48, text="BUFFETT SCORE", font=("Segoe UI", 8, "bold"), fill=GOLD)


# ============================================================================
# GUI — APPLICAZIONE PRINCIPALE
# ============================================================================
DOC_TEXT = ''
try:
    with open(os.path.join(BASE_DIR, 'docs_guida.txt'), encoding='utf-8') as _f:
        DOC_TEXT = _f.read()
except Exception:
    pass


class BuffettApp:
    def __init__(self, root):
        self.root = root
        self.result = None
        root.title(f"{APP_TITLE} v{VERSION}")
        root.configure(bg=BG)
        root.geometry("1280x820")
        root.minsize(1000, 640)

        style = ttk.Style(root)
        try:
            style.theme_use("clam")
        except Exception:
            pass
        style.configure("TNotebook", background=BG, borderwidth=0, tabmargins=[12, 8, 12, 0])
        style.configure("TNotebook.Tab", background=PANEL, foreground=MUTED,
                        padding=[20, 9], font=("Segoe UI", 10, "bold"), borderwidth=0)
        style.map("TNotebook.Tab",
                  background=[("selected", CARD2)],
                  foreground=[("selected", GOLD)])
        style.configure("Vertical.TScrollbar", background=CARD2, troughcolor=BG,
                        borderwidth=0, arrowcolor=MUTED)
        style.configure("gold.Horizontal.TProgressbar", troughcolor=PANEL,
                        background=GOLD, borderwidth=0, thickness=6)

        self._build_header()
        self._build_toolbar()
        self._build_notebook()
        self._build_doc_tab()
        self._build_statusbar()
        self._show_welcome()

    def _build_header(self):
        h = tk.Frame(self.root, bg=PANEL)
        h.pack(fill="x")
        tk.Label(h, text="⚖  BUFFETT ANALYZER", bg=PANEL, fg=GOLD, font=F_TITLE).pack(side="left", padx=20, pady=12)
        tk.Label(h, text="40 criteri quantitativi  •  12 criteri qualitativi  •  DCF Owner Earnings  •  100% automatico",
                 bg=PANEL, fg=MUTED, font=F_SUB).pack(side="left", padx=8, pady=20)

    def _build_toolbar(self):
        t = tk.Frame(self.root, bg=BG)
        t.pack(fill="x", padx=16, pady=(12, 4))
        self.btn_load = tk.Button(t, text="📂  Carica Report (qualsiasi formato)", command=self.on_load,
                                  bg=GOLD, fg="#14100a", font=("Segoe UI", 10, "bold"),
                                  bd=0, padx=18, pady=8, cursor="hand2", activebackground="#ffd066")
        self.btn_load.pack(side="left")
        self.btn_demo = tk.Button(t, text="▶  Demo", command=self.on_demo,
                                  bg=CARD2, fg=FG, font=("Segoe UI", 10, "bold"),
                                  bd=0, padx=16, pady=8, cursor="hand2", activebackground=BORDER)
        self.btn_demo.pack(side="left", padx=10)
        self.btn_html = tk.Button(t, text="🌐  Esporta HTML", command=self.on_export, state="disabled",
                                  bg=CARD2, fg=FG, font=("Segoe UI", 10, "bold"),
                                  bd=0, padx=16, pady=8, cursor="hand2", activebackground=BORDER)
        self.btn_html.pack(side="left")
        self.status_lbl = tk.Label(t, text="In attesa di un report…", bg=BG, fg=MUTED, font=F_SMALL)
        self.status_lbl.pack(side="right")
        self.prog = ttk.Progressbar(t, style="gold.Horizontal.TProgressbar", length=220, maximum=100)
        self.prog.pack(side="right", padx=12)

    def _build_notebook(self):
        self.nb = ttk.Notebook(self.root)
        self.nb.pack(fill="both", expand=True, padx=16, pady=8)
        self.tabs = {}
        for key, label in [("dash", "◈  DASHBOARD"), ("qual", "⑫  CRITERI QUALITATIVI"),
                           ("quant", "CRITERI QUANTITATIVI"), ("dcf", "∑  DCF & VALUTAZIONE"),
                           ("log", "≡  LOG ESTRAZIONE"), ("doc", "📖  DOCUMENTAZIONE")]:
            if key == "doc":
                fr = tk.Frame(self.nb, bg=BG)
                self.nb.add(fr, text=label)
                self.tabs[key] = fr
            else:
                sf = ScrollableFrame(self.nb)
                self.nb.add(sf, text=label)
                self.tabs[key] = sf.inner

    def _build_statusbar(self):
        sb = tk.Frame(self.root, bg=PANEL)
        sb.pack(fill="x", side="bottom")
        tk.Label(sb, text=f"Progetto: {BASE_DIR}   ·   v{VERSION}   ·   {datetime.date.today().strftime('%d/%m/%Y')}",
                 bg=PANEL, fg=MUTED, font=F_TINY, padx=14, pady=5).pack(side="left")

    def on_load(self):
        paths = filedialog.askopenfilenames(
            title="Seleziona uno o più report aziendali",
            filetypes=[("Report", "*.pdf *.txt *.html *.htm *.csv *.xlsx *.docx *.md"),
                       ("PDF", "*.pdf"), ("Tutti", "*.*")])
        if paths:
            paths = list(paths)
            if len(paths) == 1:
                self.load_file(paths[0])
            else:
                self.load_files(paths)

    def load_files(self, paths):
        self.btn_load.config(state="disabled")
        self.btn_demo.config(state="disabled")
        self.prog["value"] = 0
        threading.Thread(target=self._worker_multi, args=(paths,), daemon=True).start()

    def _worker_multi(self, paths):
        results = []
        for i, p in enumerate(paths):
            if self.root.winfo_exists():
                self.root.after(0, lambda i=i, p=p: self._progress(
                    int(i * 100 / len(paths)), f"Analisi {i+1}/{len(paths)}: {os.path.basename(p)}"))
            try:
                results.append(analyze_document(p))
            except Exception as e:
                print(f"{os.path.basename(p)} | ERRORE: {e}")
        if results and self.root.winfo_exists():
            self.root.after(0, lambda r=results: self._show_multi(r))

    def _show_multi(self, results):
        self._show_result(results[-1])
        if "cmp" not in self.tabs:
            sf = ScrollableFrame(self.nb)
            self.nb.add(sf, text="⚖  COMPARATIVA")
            self.tabs["cmp"] = sf.inner
        fr = self.tabs["cmp"]
        for w in fr.winfo_children():
            w.destroy()
        tk.Label(fr, text=f"CONFRONTO MULTI-REPORT ({len(results)} analisi)", bg=BG, fg=GOLD,
                 font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(16, 8))
        for r in results:
            s = r["scores"]; d = r["D"]
            roe = next((m.value for m in r["quant"] if m.code == "Q03"), None)
            mos = r["dcf"].get("mos")
            tk.Label(fr, text=f"• {r['company']} — {r['source']} · Score {s['final']:.1f}/100 · "
                              f"Qualità {s['qual']:.0f} · Numeri {s['quant']:.0f} · "
                              f"Ricavi {fmt(d.get('ricavi'), 0)} · Utile {fmt(d.get('net_income'), 0)} · "
                              f"ROE {fmt(roe, 1) if roe else 'N/D'} · MoS {fmt_pct(mos)}",
                     bg=BG, fg=FG, font=F_SMALL, anchor="w", justify="left",
                     wraplength=900).pack(fill="x", padx=24, pady=2)
        try:
            self.nb.select(self.nb.tabs()[-1])
        except Exception:
            pass

    def on_demo(self):
        ensure_project()
        self.load_file(SAMPLE_FILE)

    def on_export(self):
        if not self.result:
            return
        path = export_html(self.result)
        if messagebox.askyesno("Report esportato", f"Dashboard HTML salvata in:\n{path}\n\nAprirla nel browser?"):
            webbrowser.open("file:///" + path.replace("\\", "/"))

    def load_file(self, path):
        self.btn_load.config(state="disabled")
        self.btn_demo.config(state="disabled")
        self.prog["value"] = 0
        threading.Thread(target=self._worker, args=(path,), daemon=True).start()

    def _worker(self, path):
        def cb(p, msg):
            if self.root.winfo_exists():
                self.root.after(0, lambda p=p, msg=msg: self._progress(p, msg))
        try:
            res = analyze_document(path, cb)
            if self.root.winfo_exists():
                self.root.after(0, lambda r=res: self._show_result(r))
        except Exception as e:
            traceback.print_exc()
            if self.root.winfo_exists():
                self.root.after(0, lambda err=e: self._on_error(err))

    def _progress(self, p, msg):
        self.prog["value"] = p
        self.status_lbl.config(text=msg)

    def _on_error(self, err):
        self.btn_load.config(state="normal")
        self.btn_demo.config(state="normal")
        self.status_lbl.config(text="Errore durante l'analisi")
        messagebox.showerror("Errore analisi", str(err))

    def _build_doc_tab(self):
        if "doc" not in self.tabs:
            fr = tk.Frame(self.nb, bg=BG)
            self.nb.add(fr, text="📖  DOCUMENTAZIONE")
            self.tabs["doc"] = fr
        fr = self.tabs["doc"]
        t = tk.Text(fr, bg=CARD, fg=FG, font=F_BODY, bd=0, padx=20, pady=16, wrap="word")
        t.tag_configure("h", foreground=GOLD, font=("Segoe UI", 13, "bold"))
        t.tag_configure("s", foreground=BLUE, font=("Segoe UI", 10, "bold"))
        for line in DOC_TEXT.splitlines():
            tag = "h" if (line.startswith("BUFFETT ANALYZER") or re.match(r"^\d+\)", line)) else ("s" if line.strip().startswith("Prezzo azione") else "")
            t.insert("end", line + "\n", tag)
        t.config(state="disabled")
        sb = ttk.Scrollbar(fr, orient="vertical", command=t.yview)
        t.config(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        t.pack(side="left", fill="both", expand=True)

        def _w(e):
            if e.num == 4:
                t.yview_scroll(-1, "units")
            elif e.num == 5:
                t.yview_scroll(1, "units")
            else:
                t.yview_scroll(int(-1 * (e.delta / 120)), "units")
        t.bind("<MouseWheel>", _w)
        t.bind("<Button-4>", _w)
        t.bind("<Button-5>", _w)

    def _show_welcome(self):
        f = self.tabs["dash"]
        clear(f)
        box = tk.Frame(f, bg=BG)
        box.pack(pady=60, padx=40)
        tk.Label(box, text="🏛", bg=BG, fg=GOLD, font=("Segoe UI", 54)).pack()
        tk.Label(box, text="Benvenuto nel Buffett Analyzer", bg=BG, fg=FG, font=("Segoe UI", 18, "bold")).pack(pady=(10, 6))
        tk.Label(box, text=("1. Carica un report aziendale (bilancio PDF o TXT) oppure premi ▶ Demo.\n"
                            "2. L'estrazione dei dati è completamente automatica: nessun inserimento manuale.\n"
                            "3. Ottieni i 40 criteri quantitativi, i 12 criteri qualitativi, il DCF\n"
                            "    su Owner Earnings e il verdetto finale in stile Berkshire Hathaway.\n\n"
                            "📌 Per i PDF è richiesta la libreria 'pypdf' (pip install pypdf)."),
                 bg=BG, fg=MUTED, font=F_BODY, justify="left").pack()

    def _show_result(self, res):

        # V13 tab
        try:
            self.nb.tab(self.tabs["quant"], text=("20  CRITERI QUANTITATIVI" if res["D"].get("_bank") else "CRITERI QUANTITATIVI"))
        except Exception:
            pass
        self.result = res
        self.btn_load.config(state="normal")
        self.btn_demo.config(state="normal")
        self.btn_html.config(state="normal")
        self._build_dashboard(res)
        self._build_qual_tab(res)
        self._build_quant_tab(res)
        self._build_dcf_tab(res)
        self._build_log_tab(res)
        self.nb.select(0)

    def _build_dashboard(self, res):
        f = self.tabs["dash"]
        clear(f)
        s, dcf = res["scores"], res["dcf"]

        top = tk.Frame(f, bg=BG)
        top.pack(fill="x", padx=24, pady=(20, 8))
        gframe = tk.Frame(top, bg=CARD, highlightbackground=BORDER, highlightthickness=1)
        gframe.pack(side="left", padx=(0, 20))
        cv = tk.Canvas(gframe, width=240, height=240, bg=CARD, highlightthickness=0)
        cv.pack(padx=10, pady=10)
        draw_gauge(cv, s["final"])

        right = tk.Frame(top, bg=BG)
        right.pack(side="left", fill="both", expand=True)
        tk.Label(right, text=res["company"], bg=BG, fg=FG, font=("Segoe UI", 15, "bold"),
                 anchor="w", wraplength=700, justify="left").pack(fill="x")
        tk.Label(right, text=f"Fonte: {res['source']}   ·   {res['text_len']:,} caratteri analizzati",
                 bg=BG, fg=MUTED, font=F_TINY, anchor="w").pack(fill="x", pady=(2, 10))
        tk.Label(right, text=res["verdict"], bg=BG, fg=GOLD, font=("Segoe UI", 16, "bold"),
                 anchor="w").pack(fill="x")
        tk.Label(right, text=res["verdict_sub"], bg=BG, fg=FG, font=F_BODY, anchor="w",
                 justify="left", wraplength=700).pack(fill="x", pady=(6, 12))
        chips = tk.Frame(right, bg=BG)
        chips.pack(fill="x")
        kpi_card(chips, "PUNTEGGIO QUALITATIVO", f"{s['qual']:.1f}/100",
                 f"{s['qual_pass']} criteri superati su {s['qual_n']}").pack(side="left", padx=(0, 10))
        kpi_card(chips, "PUNTEGGIO QUANTITATIVO", f"{s['quant']:.1f}/100",
                 f"{s['quant_n']} metriche valutate").pack(side="left", padx=(0, 10))
        mos = dcf.get("mos") if dcf.get("ok") else None
        mos_col = GREEN if (mos or -99) >= 30 else (AMBER if (mos or -99) >= 20 else RED)
        kpi_card(chips, "MARGINE DI SICUREZZA", fmt_pct(mos), "sconto richiesto 20–30%", mos_col).pack(side="left")
        rub = tk.Frame(f, bg=BG)
        rub.pack(fill="x", padx=24, pady=(4, 8))
        tk.Label(rub, text="RUBRICA ANALITICA PONDERATA", bg=BG, fg=GOLD, font=F_CARD_T, anchor="w").pack(fill="x")
        for an, aw, av in s.get("q_areas", []) + s.get("k_areas", []):
            tk.Label(rub, text=f"• {an} · peso {aw*100:.0f}% · {av:.1f}/100", bg=BG, fg=FG, font=F_SMALL, anchor="w").pack(fill="x")
        tk.Label(rub, text=f"Descrittore: {descriptor_for(s['final'])}", bg=BG, fg=MUTED, font=F_SMALL, anchor="w").pack(fill="x")

        D = res["D"]
        grid = tk.Frame(f, bg=BG)
        grid.pack(fill="x", padx=24, pady=12)
        items = [
            ("RICAVI", fmt(D.get("ricavi"), 0)),
            ("UTILE NETTO", fmt(D.get("net_income"), 0)),
            ("OWNER EARNINGS", fmt((D.get("cfo") - D.get("capex")) if (D.get("cfo") is not None and D.get("capex") is not None) else None, 0)),
            ("ROIC", next((m.value_text for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")] if m.code == "Q02"), "N/D")),
            ("ROE", next((m.value_text for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")] if m.code == "Q03"), "N/D")),
            ("DEBT/EQUITY", next((m.value_text for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")] if m.code == "Q05"), "N/D")),
        ]
        if D.get("cet1") is not None:
            items.append(("CET1", fmt_pct(D["cet1"] * 100)))
        if D.get("cost_income") is not None:
            items.append(("COST/INCOME", fmt_pct(D["cost_income"] * 100)))
        for i, (lab, val) in enumerate(items):
            kpi_card(grid, lab, val).grid(row=i // 3, column=i % 3, sticky="nsew", padx=6, pady=6)
        for c in range(3):
            grid.columnconfigure(c, weight=1, uniform="a")

        if dcf.get("ok"):
            box = tk.Frame(f, bg=GOLD, padx=2, pady=2)
            box.pack(fill="x", padx=24, pady=(8, 24))
            inner = tk.Frame(box, bg=CARD)
            inner.pack(fill="x")
            txt = (f"VALORE INTRINSECO multi-scenario — Bear {fmt(dcf.get('iv_bear'),0)} · Base {fmt(dcf['iv'],0)} · Bull {fmt(dcf.get('iv_bull'),0)}"
                   + (f"   ·   per azione {fmt(dcf['iv_share'],2)}" if dcf.get("iv_share") else "")
                   + (f"   ·   prezzo {fmt(dcf['price'],2)}" if dcf.get("price") else "")
                   + (f"   ·   MoS bear/base/bull: {fmt_pct(dcf.get('mos_bear'))}/{fmt_pct(dcf.get('mos'))}/{fmt_pct(dcf.get('mos_bull'))}" if dcf.get("mos") is not None else "")
                   + (f"   ·   Tangible book {fmt(dcf.get('bv'),0)}" if dcf.get("bv") else "")
                   + (f"   ·   DDM {fmt(dcf.get('ddm'),0)}" if dcf.get("ddm") else ""))
            tk.Label(inner, text="💰 " + txt, bg=CARD, fg=GOLD, font=("Segoe UI", 11, "bold"),
                     anchor="w", padx=14, pady=12, wraplength=900, justify="left").pack(fill="x")

    def _build_qual_tab(self, res):
        f = self.tabs["qual"]
        clear(f)
        tk.Label(f, text="I 12 CRITERI QUALITATIVI DI WARREN BUFFETT", bg=BG, fg=GOLD,
                 font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(18, 4))
        tk.Label(f, text="Approvazione qualitativa: ogni criterio è verificato tramite la propria spia matematica.",
                 bg=BG, fg=MUTED, font=F_SMALL, anchor="w").pack(fill="x", padx=24, pady=(0, 12))
        for q in res["qual"]:
            if q.get("hidden") or "Non significativo" in (q.get("note") or ""): continue
            qual_card(f, q).pack(fill="x", padx=24, pady=6)
        adv = res.get("adv", [])
        if adv:
            tk.Label(f, text="PROXY QUALITATIVI AVANZATI — SPIE DI SECONDO LIVELLO", bg=BG, fg=GOLD,
                     font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(14, 4))
            for ad in adv:
                card = tk.Frame(f, bg=CARD, highlightbackground=BORDER, highlightthickness=1, padx=16, pady=12)
                card.pack(fill="x", padx=24, pady=6)
                col, lab = QSTATUS_META.get(ad["status"], (GRAY, "… N/D"))
                top = tk.Frame(card, bg=CARD); top.pack(fill="x")
                tk.Label(top, text=ad["name"], bg=CARD, fg=FG, font=("Segoe UI", 11, "bold"), anchor="w").pack(side="left")
                tk.Label(top, text=lab, bg=col, fg="#0a0f16", font=("Segoe UI", 8, "bold"), padx=8, pady=2).pack(side="right")
                tk.Label(card, text=ad["note"], bg=CARD, fg=col, font=F_SMALL, anchor="w",
                         justify="left", wraplength=620).pack(fill="x", pady=(6, 0))
        tk.Frame(f, bg=BG, height=20).pack()

    def _build_quant_tab(self, res):
        f = self.tabs["quant"]
        clear(f)
        tk.Label(f, text=("I 20 CRITERI BANCARI — METRICHE DI SETTORE (B1–B6) + VALUTAZIONE" if res["D"].get("_bank") else ("I 20 CRITERI QUANTITATIVI (FINANCIAL) — METRICHE BANCARIE + VALUTAZIONE" if res["D"].get("_bank") else ("I 20 CRITERI QUANTITATIVI (BANKING) — METRICHE DI SETTORE + VALUTAZIONE" if res["D"].get("_bank") else "I 40 CRITERI QUANTITATIVI — CORE BUFFETT (1–21) + COMPLEMENTARI (22–40)"))),
                 bg=BG, fg=GOLD, font=F_H2, anchor="w").grid(row=0, column=0, columnspan=2, sticky="ew", padx=24, pady=(18, 12))
        f.columnconfigure(0, weight=1, uniform="a")
        f.columnconfigure(1, weight=1, uniform="a")
        for i, m in enumerate([m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")]):
            metric_card(f, m).grid(row=1 + i // 2, column=i % 2, sticky="nsew", padx=(24 if i % 2 == 0 else 8, 24 if i % 2 else 8), pady=6)

    def _build_dcf_tab(self, res):
        f = self.tabs["dcf"]
        clear(f)
        dcf = res["dcf"]
        tk.Label(f, text="VALUTAZIONE DCF SU OWNER EARNINGS + MARGINE DI SICUREZZA",
                 bg=BG, fg=GOLD, font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(18, 8))
        if not dcf.get("ok"):
            tk.Label(f, text="Dati insufficienti per il DCF (serve un flusso di cassa o un utile positivo).",
                     bg=BG, fg=RED, font=F_BODY, anchor="w").pack(fill="x", padx=24)
            return
        par = tk.Frame(f, bg=BG)
        par.pack(fill="x", padx=24, pady=6)
        kpi_card(par, "FLUSSO BASE", fmt(dcf["base"], 0),
                 ("Patrimonio netto · Residual Income Model" if dcf.get("base_is_bv") else ("Owner Earnings (CFO − CapEx)" if dcf["base_is_oe"] else "FCF / Utile (fallback)"))).grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        kpi_card(par, "CRESCITA FASE 1", fmt_pct(dcf["g1"] * 100), "prudenziale, max 8%").grid(row=0, column=1, sticky="nsew", padx=6, pady=6)
        kpi_card(par, "TASSO DI SCONTO", fmt_pct(dcf["r"] * 100), "conservativo").grid(row=0, column=2, sticky="nsew", padx=6, pady=6)
        kpi_card(par, "CRESCITA PERPETUA", fmt_pct(dcf["g"] * 100), "terminal value").grid(row=0, column=3, sticky="nsew", padx=6, pady=6)
        for c in range(4):
            par.columnconfigure(c, weight=1, uniform="a")

        rows = tk.Frame(f, bg=BG)
        rows.pack(fill="x", padx=24, pady=10)
        tk.Label(rows, text="Flussi attualizzati (10 anni):", bg=BG, fg=MUTED, font=F_SMALL, anchor="w").pack(fill="x")
        flow_txt = "   ".join(f"Anno {i+1}: {fmt(v,0)}" for i, v in enumerate(dcf["flows"][:5]))
        flow_txt2 = "   ".join(f"Anno {i+6}: {fmt(v,0)}" for i, v in enumerate(dcf["flows"][5:]))
        tk.Label(rows, text=flow_txt, bg=BG, fg=FG, font=F_MONO, anchor="w").pack(fill="x", pady=(4, 0))
        tk.Label(rows, text=flow_txt2, bg=BG, fg=FG, font=F_MONO, anchor="w").pack(fill="x")
        tk.Label(rows, text=f"Terminal Value attualizzato: {fmt(dcf['tv'] / (1 + dcf['r']) ** dcf['n'], 0)}",
                 bg=BG, fg=FG, font=F_MONO, anchor="w").pack(fill="x", pady=(4, 0))
        if dcf.get("ddm"):
            tk.Label(rows, text=f"Cross-check DDM (Gordon growth sui dividendi): {fmt(dcf['ddm'], 0)}",
                     bg=BG, fg=MUTED, font=F_MONO, anchor="w").pack(fill="x", pady=(4, 0))

        summ = tk.Frame(f, bg=BG)
        summ.pack(fill="x", padx=24, pady=12)
        kpi_card(summ, "VALORE INTRINSECO", fmt(dcf["iv"], 0),
                 f"per azione: {fmt(dcf.get('iv_share'),2)}" if dcf.get("iv_share") else "", GOLD).grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        kpi_card(summ, "PREZZO / MARKET CAP", fmt(dcf.get("price"), 2) if dcf.get("price") else fmt(dcf.get("mcap"), 0),
                 "valore di mercato" if dcf.get("mcap") else "").grid(row=0, column=1, sticky="nsew", padx=6, pady=6)
        mos = dcf.get("mos")
        mos_col = GREEN if (mos or -99) >= 30 else (AMBER if (mos or -99) >= 20 else RED)
        kpi_card(summ, "MARGINE DI SICUREZZA", fmt_pct(mos),
                 "Buffett richiede ≥ 20–30%", mos_col).grid(row=0, column=2, sticky="nsew", padx=6, pady=6)
        for c in range(3):
            summ.columnconfigure(c, weight=1, uniform="a")

        tk.Label(f, text=res["verdict_sub"], bg=BG, fg=FG, font=F_BODY, anchor="w",
                 wraplength=860, justify="left").pack(fill="x", padx=30, pady=(4, 24))

    def _build_log_tab(self, res):
        f = self.tabs["log"]
        clear(f)
        t = tk.Text(f, bg=CARD, fg=FG, font=F_MONO, bd=0, padx=14, pady=12, wrap="none")
        t.tag_configure("ok", foreground=GREEN)
        t.tag_configure("warn", foreground=AMBER)
        t.tag_configure("nd", foreground=GRAY)
        t.insert("end", f"ANALISI AUTOMATICA — {res['source']}\n", "ok")
        t.insert("end", f"Azienda: {res['company']}\n")
        t.insert("end", f"Caratteri estratti: {res['text_len']:,}\n")
        t.insert("end", f"Grandezze trovate: {sum(1 for k in res['D'] if k != '_unit')}\n\n")
        for lvl, msg in res["log"]:
            prefix = "✔ " if lvl == "ok" else ("⚠ " if lvl == "warn" else "· ")
            t.insert("end", prefix + msg + "\n", lvl)
        t.config(state="disabled")
        t.pack(fill="both", expand=True, padx=24, pady=18)


# ============================================================================
# EXPORT HTML
# ============================================================================
HTML_CSS = """
:root{--bg:#0b0e14;--card:#161c28;--border:#232c3f;--fg:#e8ecf4;--muted:#8a94ab;
--gold:#f2b632;--green:#2ecc71;--amber:#f39c12;--red:#ff5c5c;--gray:#5a6478;}
*{box-sizing:border-box;margin:0;padding:0}
body{background:var(--bg);color:var(--fg);font-family:'Segoe UI',Arial,sans-serif;padding:32px}
.hero{background:linear-gradient(135deg,#161c28,#1b2333);border:1px solid var(--border);
border-radius:16px;padding:32px;margin-bottom:24px}
.hero h1{color:var(--gold);font-size:26px;margin-bottom:6px}
.hero .sub{color:var(--muted);font-size:13px}
.score{font-size:52px;font-weight:800;color:var(--gold);margin-top:12px}
.verdict{font-size:18px;margin-top:8px;font-weight:600}
.verdict-sub{color:var(--muted);margin-top:6px;max-width:900px}
h2{color:var(--gold);margin:28px 0 14px;font-size:18px;letter-spacing:.5px}
.grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(340px,1fr));gap:12px}
.card{background:var(--card);border:1px solid var(--border);border-radius:12px;padding:16px 18px;
border-left:4px solid var(--gray)}
.card.pass{border-left-color:var(--green)} .card.warn,.card.partial{border-left-color:var(--amber)}
.card.fail{border-left-color:var(--red)}
.card .name{font-weight:700;font-size:14px}
.card .tag{color:var(--muted);font-size:11px;margin:2px 0 8px}
.card .value{font-size:22px;font-weight:800}
.card .thr{color:var(--muted);font-size:11px;margin-top:6px}
.card .det{color:#4da3ff;font-size:11px;margin-top:4px}
.card .note{color:var(--muted);font-size:12px;margin-top:6px}
.pill{display:inline-block;padding:2px 10px;border-radius:10px;font-size:11px;font-weight:700;
color:#0a0f16;float:right}
.pill.pass{background:var(--green)} .pill.warn,.pill.partial{background:var(--amber)}
.pill.fail{background:var(--red)} .pill.nd{background:var(--gray);color:#dfe4ee}
footer{color:var(--muted);font-size:11px;margin-top:36px;text-align:center}
"""


def export_html(res):
    s, dcf = res["scores"], res["dcf"]
    cards_q = []
    for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")]:
        if getattr(m, "hidden", False) or "Non significativo" in (m.detail or ""): continue
        if res["D"].get("_bank") and m.status == "nd": continue
        _, lab = STATUS_META.get(m.status, (GRAY, "… N/D"))
        safe_name = html.escape(m.name)
        safe_tag = html.escape(m.tag)
        safe_val = html.escape(m.value_text)
        safe_thr = html.escape(m.threshold)
        safe_det = html.escape(m.detail) if m.detail else ""
        cards_q.append(
            f'<div class="card {m.status}"><span class="pill {m.status}">{lab}</span>'
            f'<div class="name">{m.code} · {safe_name}</div><div class="tag">{safe_tag}'
            f'{" · CORE BUFFETT" if m.core else ""}</div>'
            f'<div class="value">{safe_val}</div>'
            f'<div class="thr">Soglia: {safe_thr}</div>'
            + (f'<div class="det">{safe_det}</div>' if safe_det else "") + "</div>")
    cards_k = []
    for q in res["qual"]:
        if q.get("hidden"): continue
        _, lab = QSTATUS_META.get(q["status"], (GRAY, "… N/D"))
        safe_title = html.escape(q['title'])
        safe_desc = html.escape(q['desc'])
        safe_formula = html.escape(q['formula'])
        safe_note = html.escape(q['note'])
        cards_k.append(
            f'<div class="card {q["status"]}"><span class="pill {q["status"]}">{lab}</span>'
            f'<div class="name">{q["n"]:02d} · {safe_title}</div>'
            f'<div class="note">{safe_desc}</div>'
            f'<div class="det">📐 {safe_formula}</div>'
            f'<div class="thr">→ {safe_note}</div></div>')
    mos_txt = fmt_pct(dcf.get("mos")) if dcf.get("ok") else "N/D"
    adv_html = ""
    if res.get("adv"):
        adv_html = "<h2>PROXY QUALITATIVI AVANZATI</h2><div class='grid'>" + "".join(
            f"<div class='card {ad['status']}'><span class='pill {ad['status']}'>"
            + html.escape(QSTATUS_META.get(ad["status"], (GRAY, "N/D"))[1]) + "</span>"
            + f"<div class='name'>{html.escape(ad['name'])}</div><div class='thr'>{html.escape(ad['note'])}</div></div>"
            for ad in res["adv"]) + "</div>"

    html_code = ("<!DOCTYPE html><html lang='it'><head><meta charset='utf-8'>"
            f"<title>Buffett Analyzer — {html.escape(res['company'])}</title><style>{HTML_CSS}</style></head><body>"
            f"<div class='hero'><h1>⚖ BUFFETT ANALYZER</h1>"
            f"<div class='sub'>{html.escape(res['company'])} · Fonte: {html.escape(res['source'])} · "
            f"{datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}</div>"
            f"<div class='score'>{s['final']:.1f} / 100</div>"
            f"<div class='verdict'>{html.escape(res['verdict'])}</div>"
            f"<div class='verdict-sub'>{html.escape(res['verdict_sub'])}</div>"
            f"<div class='sub' style='margin-top:12px'>Qualitativo {s['qual']:.0f}/100 · "
            f"Quantitativo {s['quant']:.0f}/100 · Margine di sicurezza {mos_txt}</div></div>"
            "<h2>⑫ CRITERI QUALITATIVI</h2><div class='grid'>" + "".join(cards_k) + "</div>"
            "<h2>40 CRITERI QUANTITATIVI</h2><div class='grid'>" + "".join(cards_q) + "</div>"
            "<footer>Generato automaticamente da Buffett Analyzer v" + VERSION + " · Metodo Berkshire Hathaway"
            " · Nessuna consulenza finanziaria</footer></body></html>")
    slug = re.sub(r"[^A-Za-z0-9]+", "_", res["company"])[:40] or "report"
    path = os.path.join(OUTPUT_DIR, f"{slug}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html_code)
    return path


# ============================================================================
# MAIN
# ============================================================================

# ============================================================================
# UNIVERSAL LOADER V2 — legge qualsiasi formato di report reale
# (sniffing magic-bytes + fallback a catena; tabelle → righe pipe)
# ============================================================================
import zipfile as _zipfile
import io as _io


def _sniff_kind(data):
    if data[:5] == b"%PDF-": return "pdf"
    if data[:4] == b"PK\x03\x04": return "zip"
    if data[:5] == b"{\\rtf": return "rtf"
    head = data[:2000].lstrip().lower()
    if head.startswith(b"<?xml") or head.startswith(b"<!doctype") or head.startswith(b"<html") or head.startswith(b"<"):
        return "html"
    return "text"


def _xml_to_text(xml):
    xml = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", xml)
    xml = re.sub(r"(?is)</(td|th)>", " | ", xml)
    xml = re.sub(r"(?is)</(p|div|tr|table|row|li|h[1-6]|text)>", "\n", xml)
    xml = re.sub(r"(?s)<[^>]+>", " ", xml)
    return html.unescape(re.sub(r"[ \t]+", " ", xml))


def _docx_via_python_docx(path):
    """v2.1: DOCX letto con python-docx (paragrafi + tabelle come righe pulite)."""
    try:
        import docx
    except Exception:
        return ""
    try:
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                cells = []
                for c in row.cells:
                    txt = c.text.strip()
                    if txt and txt not in cells:
                        cells.append(txt)
                if cells:
                    parts.append("  ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _docx_text(z):
    xml = z.read("word/document.xml").decode("utf-8", "ignore")
    xml = re.sub(r"(?is)</w:(tc|th)>", " | ", xml)
    xml = re.sub(r"(?is)</w:(tr|p|tbl)>", "\n", xml)
    xml = re.sub(r"(?s)<[^>]+>", " ", xml)
    return html.unescape(re.sub(r"[ \t]+", " ", xml))


def _xlsx_text(z):
    try:
        ss = []
        if "xl/sharedStrings.xml" in z.namelist():
            sx = z.read("xl/sharedStrings.xml").decode("utf-8", "ignore")
            ss = [html.unescape(re.sub(r"(?s)<[^>]+>", "", s))
                  for s in re.findall(r"(?s)<si>.*?</si>", sx)]
        out = []
        sheets = sorted(n for n in z.namelist() if re.match(r"xl/worksheets/sheet\d+\.xml", n))
        for sh in sheets:
            sx = z.read(sh).decode("utf-8", "ignore")
            for row in re.findall(r"(?s)<row[^>]*>.*?</row>", sx):
                cells = []
                for c in re.findall(r"(?s)<c[^>]*>.*?</c>|<c[^>]*/>", row):
                    t = re.search(r't="([^"]+)"', c)
                    v = re.search(r"(?s)<v>(.*?)</v>", c)
                    ti = re.search(r"(?s)<t[^>]*>(.*?)</t>", c)
                    val = ""
                    if t and t.group(1) == "s" and v:
                        try: val = ss[int(v.group(1))]
                        except Exception: val = ""
                    elif t and t.group(1) == "inlineStr" and ti:
                        val = html.unescape(ti.group(1))
                    elif v:
                        val = v.group(1)
                    if val.strip():
                        cells.append(val.strip())
                if cells:
                    out.append("  ".join(cells))
        return "\n".join(out)
    except Exception:
        return ""


def _rtf_text(data):
    t = data.decode("latin-1", "ignore")
    t = re.sub(r"(?s)\{\*[^}]*}", " ", t)
    t = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: chr(int(m.group(1), 16)), t)
    t = re.sub(r"\\par\b", "\n", t)
    t = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", t)
    t = re.sub(r"[{}]", "", t)
    return t


def _json_text(data):
    import json
    obj = json.loads(data.decode("utf-8", "ignore"))
    lines = []
    def walk(o, pre=""):
        if isinstance(o, dict):
            for k, v in o.items():
                walk(v, (pre + " " + str(k)).strip())
        elif isinstance(o, list):
            for i in o:
                walk(i, pre)
        else:
            lines.append(f"{pre}: {o}")
    walk(obj)
    return "\n".join(lines)


def _ocr_pdf_bytes(data):
    try:
        from pdf2image import convert_from_path
        import pytesseract, tempfile
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
            tf.write(data); tmp = tf.name
        try:
            return "\n".join(pytesseract.image_to_string(p, lang="eng+ita") for p in convert_from_path(tmp, dpi=200))
        except Exception:
            return "\n".join(pytesseract.image_to_string(p) for p in convert_from_path(tmp, dpi=200))
    except Exception:
        return ""


def extract_text_from_file(path):
    """Loader universale v2: riconosce il formato dal contenuto, non solo dall'estensione."""
    ext = os.path.splitext(path)[1].lower()
    with open(path, "rb") as f:
        data = f.read()
    kind = _sniff_kind(data)

    if ext == ".pdf" or kind == "pdf":
        text = ""
        if PdfReader is not None:
            try:
                r = PdfReader(_io.BytesIO(data))
                text = "\n".join((p.extract_text() or "") for p in r.pages)
            except Exception:
                text = ""
        if not text.strip():
            text = _ocr_pdf_bytes(data)
        if not text.strip():
            raise ValueError("PDF senza testo estraibile e OCR non disponibile (pip install pypdf pytesseract pdf2image).")
        return text

    if kind == "zip":
        with _zipfile.ZipFile(_io.BytesIO(data)) as z:
            names = set(z.namelist())
            if "word/document.xml" in names:
                t = _docx_via_python_docx(path)
                if t.strip():
                    return t
                return _docx_text(z)
            if "xl/workbook.xml" in names:
                t = _xlsx_text(z)
                if t.strip():
                    return t
                raise ValueError("XLSX non leggibile: pip install openpyxl")
            if "content.xml" in names:
                return _xml_to_text(z.read("content.xml").decode("utf-8", "ignore"))
            outs = []
            for n in sorted(names):
                if n.lower().endswith((".txt", ".xml", ".html", ".htm")):
                    try:
                        outs.append(_xml_to_text(z.read(n).decode("utf-8", "ignore")))
                    except Exception:
                        pass
            if outs:
                return "\n".join(outs)
            raise ValueError("Archivio ZIP senza contenuti testuali leggibili.")

    if ext == ".rtf" or kind == "rtf":
        return _rtf_text(data)

    if ext in (".html", ".htm", ".xhtml", ".xml", ".xbrl", ".sgm", ".sgml") or kind == "html":
        return _xml_to_text(data.decode("utf-8", "ignore"))

    if ext == ".json":
        try:
            return _json_text(data)
        except Exception:
            pass

    if ext in (".csv", ".tsv"):
        import csv
        raw = data.decode("utf-8", "ignore")
        out = []
        for row in csv.reader(_io.StringIO(raw), delimiter="\t" if ext == ".tsv" else ","):
            cells = [c.strip() for c in row if c and c.strip()]
            if cells:
                out.append("  ".join(cells))
        return "\n".join(out)

    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    raise ValueError("Formato file non leggibile")



def main():
    ensure_project()
    root = tk.Tk()
    app = BuffettApp(root)
    if len(sys.argv) > 1 and os.path.exists(sys.argv[1]):
        root.after(400, lambda: app.load_file(os.path.abspath(sys.argv[1])))
    root.mainloop()


# ============================ V3 "IMPECCABILE" ============================
import json as _json3

_SEG_DEFS = [
    ("Americas", r"\bAmericas\b"), ("Europe", r"\bEurope\b"),
    ("Greater China", r"\bGreater China\b"), ("Japan", r"\bJapan\b"),
    ("Rest of Asia Pacific", r"\bRest of Asia Pacific\b"),
    ("Productivity & Business Processes", r"\bProductivity and Business Processes\b"),
    ("Intelligent Cloud", r"\bIntelligent Cloud\b"),
    ("More Personal Computing", r"\bMore Personal Computing\b"),
]

def compute_segments(text):
    out = []
    for label, pat in _SEG_DEFS:
        m = re.search(pat + r"[\s\S]{0,160}?[\$]?\s*([\d][\d,]{4,})", text)
        if m:
            try:
                rev = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1))
            except Exception:
                rev = None
            if rev and rev > 1000:
                out.append({"label": label, "rev": rev})
    tot = sum(s["rev"] for s in out) if out else 0
    for s in out:
        s["share"] = (s["rev"] / tot * 100) if tot else None
    return out

def compute_adv_v3(text, D, S):
    adv = []
    eq = D.get("equity") or 0
    ni = D.get("net_income")
    utb = nopa = None
    m = re.search(r"unrecognized tax benefits (?:was|were)\s*\$?\s*([\d.,]+)\s*billion", text, re.I)
    if m: utb = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1)) * 1000
    m = re.search(r"additional tax payment of\s*\$?\s*([\d.,]+)\s*billion", text, re.I)
    if m: nopa = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1)) * 1000
    if (utb or nopa) and eq:
        val = max(x for x in (utb, nopa) if x)
        r_ = val / eq * 100
        adv.append({"name": "A10 · Posizioni fiscali incerte vs patrimonio",
                    "status": "pass" if r_ < 10 else ("warn" if r_ < 25 else "fail"),
                    "note": f"UTB/NOPA ≈ {fmt(val,0)} ({fmt(r_,1)}% equity): " + ("contenuto" if r_ < 10 else "rilevante: leggere Note imposte")})
    cams = len(re.findall(r"How (?:We Addressed|the Critical Audit Matter Was Addressed)", text))
    if cams:
        adv.append({"name": "A11 · Critical Audit Matters", "status": "pass" if cams <= 1 else "warn",
                    "note": f"{cams} CAM: " + ("contabilità a basso giudizio" if cams <= 1 else "aree con forte giudizio/stime (revenue recognition, tasse): leggere le Note")})
    m = re.search(r"auditor since\s+((?:19|20)\d{2})", text, re.I)
    if m:
        yr = int(m.group(1)); yrs = 2026 - yr
        adv.append({"name": "A12 · Tenure revisore", "status": "pass" if yrs <= 20 else "warn",
                    "note": f"Revisore dal {yr} ({yrs} anni): " + ("indipendenza recente" if yrs <= 20 else "tenure lunga: verificare rotazione/controlli")})
    comm = leases = None
    m = re.search(r"manufacturing purchase obligations of\s*\$?\s*([\d.,]+)\s*billion", text, re.I)
    if m: comm = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1)) * 1000
    leases = D.get("lease_liab")
    totc = (comm or 0) + (leases or 0)
    if totc and eq:
        r_ = totc / eq * 100
        adv.append({"name": "A13 · Impegni contrattuali vs patrimonio",
                    "status": "pass" if r_ < 50 else ("warn" if r_ < 120 else "fail"),
                    "note": f"Impegni (purchase obligations + lease) ≈ {fmt(totc,0)} = {fmt(r_,0)}% equity"})
    ca, cl = D.get("current_assets"), D.get("current_liab")
    if ca and cl and not D.get("_bank") and (ca - cl) < 0:
        adv.append({"name": "A14 · Working capital negativo", "status": "pass",
                    "note": f"WC = {fmt(ca - cl,0)}: il business si finanzia con fornitori/clienti (segnale franchise)"})
    da, cx = D.get("depreciation"), D.get("capex")
    if da and cx:
        r_ = cx / da
        adv.append({"name": "A15 · CapEx / Ammortamenti", "status": "pass" if r_ <= 1.5 else ("warn" if r_ <= 2.5 else "fail"),
                    "note": f"CapEx {fmt(cx,0)} vs D&A {fmt(da,0)} = {fmt(r_,2)}x: " + ("mantenimento leggero" if r_ <= 1.5 else "forte reinvestimento (capex-intensive)")})
    m = re.search(r"10% decrease in foreign exchange rates[\s\S]{0,80}?\(([\d,]+)", text, re.I)
    if m and ni:
        fx = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1))
        r_ = fx / ni * 100
        adv.append({"name": "A16 · Sensibilità FX (shock -10%)", "status": "pass" if r_ < 10 else "warn",
                    "note": f"Impatto -10% FX ≈ {fmt(fx,0)} ({fmt(r_,1)}% utile): esposizione " + ("limitata" if r_ < 10 else "rilevante")})
    m = re.search(r"(?:Stock-based|Share-based) compensation expense\s*\$?\s*([\d,]+)", text, re.I)
    if m and ni:
        sbc = (lambda g: float(g.replace(",", "")) if g.strip() else 0.0)(m.group(1))
        r_ = sbc / ni * 100
        adv.append({"name": "A17 · Stock-based compensation / utile", "status": "pass" if r_ < 10 else ("warn" if r_ < 20 else "fail"),
                    "note": f"SBC = {fmt(sbc,0)} = {fmt(r_,1)}% utile: " + ("diluzione contenuta" if r_ < 10 else "diluzione significativa")})
    return adv

def _implied_growth(dcf):
    if not dcf.get("ok") or not dcf.get("mcap") or not dcf.get("base"):
        return None
    mcap, base = dcf["mcap"], dcf["base"]
    r, g, n = dcf["r"], dcf["g"], dcf["n"]
    def val(g1):
        pv = sum(base * (1 + g1) ** t / (1 + r) ** t for t in range(1, n + 1))
        return pv + base * (1 + g1) ** n * (1 + g) / (r - g) / (1 + r) ** n
    lo, hi = -0.10, 0.30
    if val(lo) > mcap: return lo
    if val(hi) < mcap: return hi
    for _ in range(60):
        mid = (lo + hi) / 2
        if val(mid) < mcap: lo = mid
        else: hi = mid
    return (lo + hi) / 2

_hist_path = os.path.join(OUTPUT_DIR, "history.json")
def _load_history():
    try:
        with open(_hist_path, encoding="utf-8") as f: return _json3.load(f)
    except Exception: return []
def _save_history(e):
    h = _load_history(); h.append(e)
    try:
        with open(_hist_path, "w", encoding="utf-8") as f: _json3.dump(h, f, ensure_ascii=False)
    except Exception: pass

_analyze_prev = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_prev(path, cb)
    text = ""
    try:
        with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
            text = f.read()
    except Exception: pass
    if text:
        res["segments"] = compute_segments(text)
        names = {a["name"] for a in res.get("adv", [])}
        try:
            _adv = list(compute_adv_v3(text, res["D"], res["S"]))
        except Exception as _e:
            print('[adv] salto proxy:', _e)
            _adv = []
        for a in _adv:
            if a["name"] not in names:
                res.setdefault("adv", []).append(a)
    res["implied_g"] = _implied_growth(res["dcf"])
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    prev = None
    for e in reversed(_load_history()):
        if e.get("company") == res["company"]:
            prev = e; break
    res["prev"], res["ts"] = prev, ts
    _save_history({"company": res["company"], "ts": ts, "final": res["scores"]["final"], "mos": res["dcf"].get("mos")})
    return res

_show_result_prev = BuffettApp._show_result
def _show_result_v3(self, res):
    _show_result_prev(self, res)
    if "seg" not in self.tabs:
        sf = ScrollableFrame(self.nb)
        self.nb.add(sf, text="🗺  APPROFONDIMENTI")
        self.tabs["seg"] = sf.inner
    fr = self.tabs["seg"]
    for w in fr.winfo_children(): w.destroy()
    tk.Label(fr, text="SEGMENTI DI BUSINESS", bg=BG, fg=GOLD, font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(16, 6))
    segs = res.get("segments", [])
    if segs:
        for s in segs:
            tk.Label(fr, text=f"• {s['label']}: ricavi {fmt(s['rev'],0)}" + (f" · {fmt(s['share'],1)}% del totale" if s.get("share") else ""),
                     bg=BG, fg=FG, font=F_SMALL, anchor="w").pack(fill="x", padx=24)
    else:
        tk.Label(fr, text="Segmenti non rilevati nel documento.", bg=BG, fg=MUTED, font=F_SMALL, anchor="w").pack(fill="x", padx=24)
    tk.Label(fr, text="CONTROLLI AVANZATI (A10–A17)", bg=BG, fg=GOLD, font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(16, 6))
    for a in res.get("adv", []):
        col, lab = QSTATUS_META.get(a["status"], (GRAY, "… N/D"))
        tk.Label(fr, text=f"[{lab}] {a['name']}: {a['note']}", bg=BG, fg=col, font=F_SMALL,
                 anchor="w", justify="left", wraplength=900).pack(fill="x", padx=24, pady=2)
    if res.get("implied_g") is not None:
        cagr = res["dcf"].get("cagr_used")
        tk.Label(fr, text="REVERSE DCF", bg=BG, fg=GOLD, font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(16, 6))
        msg = (f"Il prezzo di mercato sconta una crescita implicita ≈ {fmt(res['implied_g']*100,1)}%/anno "
               f"(vs crescita storica utili {fmt((cagr or 0)*100,1)}%): ")
        msg += ("il mercato si aspetta molto più dello storico → pretesa di prezzo alta."
                if (cagr is not None and res["implied_g"] > cagr + 0.03) else
                "crescita attesa in linea o sotto lo storico → prezzo meno esigente.")
        tk.Label(fr, text=msg, bg=BG, fg=FG, font=F_SMALL, anchor="w", justify="left", wraplength=900).pack(fill="x", padx=24)
    if res.get("prev"):
        p = res["prev"]; d = res["scores"]["final"] - p["final"]
        tk.Label(fr, text="STORICO ANALISI", bg=BG, fg=GOLD, font=F_H2, anchor="w").pack(fill="x", padx=24, pady=(16, 6))
        tk.Label(fr, text=f"Precedente ({p['ts']}): {fmt(p['final'],1)}/100 → oggi {fmt(res['scores']['final'],1)}/100 ({'+' if d >= 0 else ''}{fmt(d,1)} punti)",
                 bg=BG, fg=FG, font=F_SMALL, anchor="w").pack(fill="x", padx=24)

BuffettApp._show_result = _show_result_v3


# ============================ V5: TABELLE STRUTTURATE DA PDF ============================
# Strategia 2 (dynamic chunking sulle intestazioni) + Strategia 3 (tabelle -> righe strutturate)
_STMT_HEADS = ["stato patrimoniale", "conto economico", "prospetto della situazione",
               "prospetto di conto economico", "prospetto dei flussi", "situazione patrimoniale",
               "consolidated statements of operations", "consolidated balance sheets",
               "consolidated statements of cash flows", "consolidated statements of income",
               "statements of operations", "statements of cash flows", "balance sheets",
               "income statements", "statement of financial position"]

def _structured_pdf_tables(path):
    try:
        import pdfplumber
    except Exception:
        return ""
    try:
        from pypdf import PdfReader as _R
    except Exception:
        try:
            from PyPDF2 import PdfReader as _R
        except Exception:
            return ""
    try:
        rdr = _R(path)
        idxs = []
        for i, p in enumerate(rdr.pages):
            t = (p.extract_text() or "").lower()
            if any(h in t for h in _STMT_HEADS):
                idxs.append(i)
            if len(idxs) >= 60:
                break
        if not idxs:
            return ""
        out = []
        with pdfplumber.open(path) as pdf:
            for i in idxs:
                try:
                    tabs = pdf.pages[i].extract_tables()
                except Exception:
                    continue
                for tb in tabs or []:
                    for row in tb:
                        cells = [str(c).replace("\n", " ").strip() for c in row if c and str(c).strip()]
                        if len(cells) >= 2 and any(ch.isdigit() for ch in cells[-1]):
                            out.append(" | ".join(cells))
        return "\n".join(out)
    except Exception:
        return ""

_extract_prev_v5 = extract_text_from_file
def extract_text_from_file(path, *a, **k):
    txt = _extract_prev_v5(path, *a, **k)
    try:
        with open(path, "rb") as f:
            is_pdf = f.read(5) == b"%PDF-"
    except Exception:
        is_pdf = False
    if is_pdf:
        extra = _structured_pdf_tables(path)
        if extra:
            txt = txt + "\n" + extra
    return txt


# ============================ V7: MARKET DATA (yfinance) ============================
_TICKER_MAP = {
    "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
    "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK", "berkshire": "BRK-B",
    "jpmorgan": "JPM", "bank of america": "BAC", "wells fargo": "WFC",
}

def _yf_market(company):
    try:
        import yfinance as yf
    except Exception:
        return None, None
    tick = None
    try:
        with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
            txt = f.read()
        m = re.search(r"\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", txt)
        if m: tick = m.group(1)
    except Exception:
        pass
    if not tick:
        low = (company or "").lower()
        for k, v in _TICKER_MAP.items():
            if k in low:
                tick = v; break
    if not tick:
        return None, None
    try:
        info = yf.Ticker(tick).info
        price = info.get("currentPrice") or info.get("regularMarketPrice")
        shares = info.get("sharesOutstanding")
        return (float(price) if price else None), (float(shares) if shares else None)
    except Exception:
        return None, None

_analyze_v7base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v7base(path, cb)
    D = res["D"]
    if D.get("price") is None or D.get("shares") is None:
        price, shares_yf = _yf_market(res.get("company", ""))
        if price and D.get("price") is None:
            D["price"] = price
            res["log"].append(("ok", f"price         = {price:.2f} (Yahoo Finance)"))
        if shares_yf and D.get("shares") is None:
            D["shares"] = shares_yf / 1e6
            res["log"].append(("ok", f"shares        = {shares_yf/1e6:,.0f} M (Yahoo Finance)"))
    return res


# ============================ V8 bank-fix ============================
BANK_HIDE_Q = {"Q01","Q04","Q05","Q11","Q13","Q14","Q24","Q25","Q31"}

def _auto_ratio(num, den, lo, hi):
    if not num or not den: return None
    for mlt in (1.0, 1000.0, 0.001):
        r = (num * mlt) / den
        if lo <= r <= hi: return r
    return num / den

def _yf_price_bank(company):
    try:
        import yfinance as yf
    except Exception:
        return None, None
    tick = None
    low = (company or "").lower()
    for k, v in _TICKER_MAP.items():
        if k in low: tick = v; break
    if not tick:
        m = re.search(r"\b(?:BIT|MIL)[:\s]+([A-Z0-9.\-]{1,6})\b", company or "")
        if m: tick = m.group(1)
    if not tick: return None, None
    try:
        t = yf.Ticker(tick)
        fi = t.fast_info
        price = None
        for k in ("currentPrice","regularMarketPrice","previousClose"):
            try:
                v = fi[k]
                if v: price = float(v); break
            except Exception: pass
        if price is None:
            price = float(t.info.get("currentPrice") or t.info.get("regularMarketPrice") or 0) or None
        shares = None
        try: shares = float(fi["sharesOutstanding"])
        except Exception:
            try: shares = float(t.info.get("sharesOutstanding") or 0) or None
            except Exception: shares = None
        return price, shares
    except Exception:
        return None, None

_analyze_v8base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v8base(path, cb)
    D = res["D"]
    if D.get("_bank"):
        ni, eq, ta = D.get("net_income"), D.get("equity"), D.get("total_assets")
        if D.get("price") is None:
            p, s_yf = _yf_price_bank(res.get("company",""))
            if p:
                D["price"] = p
                if s_yf and not D.get("shares"): D["shares"] = s_yf/1e6
                if D.get("shares"): D["market_cap"] = p*D["shares"]
                res["log"].append(("ok", f"price         = {p:.2f} (Yahoo Finance)"))
        for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")]:
            c = m.code
            if c in BANK_HIDE_Q or c == "Q03":
                m.status = "nd"; m.hidden = True
            if c == "B1":
                r_ = _auto_ratio(ni, eq, 0.01, 0.80)
                if r_ is not None:
                    m.value = r_*100; m.value_text = fmt_pct(r_*100)
                    m.status = "pass" if r_*100 > 12 else ("warn" if r_*100 > 10 else "fail")
            if c == "B2":
                r_ = _auto_ratio(ni, ta, 0.001, 0.06)
                if r_ is not None:
                    m.value = r_*100; m.value_text = fmt_pct(r_*100)
                    m.status = "pass" if r_*100 > 1.0 else ("warn" if r_*100 > 0.7 else "fail")
            if c == "B6":
                pct = D.get("npl_ratio_pct")
                npl_amt, loans = D.get("npl_amount"), D.get("loans")
                if pct is None and npl_amt and loans and npl_amt < loans:
                    pct = npl_amt/loans*100
                if pct is not None and pct > 25: pct = None
                if pct is not None:
                    m.value = pct; m.value_text = fmt_pct(pct)
                    m.status = "pass" if pct < 3 else ("warn" if pct < 5 else "fail")
        for q in res["qual"]:
            if q["n"] in (4,5,9,12):
                q["status"] = "nd"; q["hidden"] = True
            if q["n"] == 8:
                roa = _auto_ratio(ni, ta, 0.001, 0.06)
                if roa is not None:
                    q["status"] = "pass" if roa*100 > 1.0 else ("partial" if roa*100 > 0.7 else "fail")
                    q["formula"] = "ROA banche > 1,0% (non 10%)"
                    q["note"] = f"ROA = {fmt(roa*100,1)}% vs soglia bancaria 1,0%"
        res["scores"] = compute_scores(res["quant"], res["qual"])
    return res

# ============================ V9 bank-final ============================
def _bank_price(company):
    tm = globals().get("_TICKER_MAP", {}) or {
        "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
        "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK"}
    try:
        import yfinance as yf
        tick = None
        low = (company or "").lower()
        for k, v in tm.items():
            if k in low: tick = v; break
        if not tick:
            mm = re.search(r"\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", company or "")
            if mm: tick = mm.group(1)
        if tick:
            t = yf.Ticker(tick)
            for attr in ("fast_info", "info"):
                try:
                    info = getattr(t, attr)
                    for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                        try:
                            v = info[k]
                            if v: return float(v), f"Yahoo Finance ({tick})"
                        except Exception: pass
                except Exception: pass
    except Exception:
        pass
    for p in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
        try:
            txt = open(p, encoding="utf-8").read()
            mm = re.search(r"([\d.,]+)", txt)
            if mm:
                s = mm.group(1)
                v = float(s.replace(".", "").replace(",", ".")) if "," in s else float(s)
                return v, "prezzo.txt"
        except Exception: pass
    return None, None

_analyze_v9base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v9base(path, cb)
    D = res["D"]
    if D.get("_bank"):
        if D.get("price") is None:
            p, srcname = _bank_price(res.get("company", ""))
            if p:
                D["price"] = p
                res["log"].append(("ok", f"price         = {p:.2f} ({srcname})"))
                if D.get("shares"): D["market_cap"] = p * D["shares"]
        rep = D.get("roe_reported") or D.get("rote")
        if rep is not None and 0.05 <= rep <= 0.60:
            for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")]:
                if m.code == "B1":
                    m.value = rep * 100; m.value_text = fmt_pct(rep * 100)
                    m.status = "pass" if rep * 100 > 12 else ("warn" if rep * 100 > 10 else "fail")
                    m.detail = f"ROE/RoTE dichiarato nel report: {fmt_pct(rep * 100)}"
        repa = D.get("roa_reported")
        if repa is not None and 0.003 <= repa <= 0.05:
            for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")]:
                if m.code == "B2":
                    m.value = repa * 100; m.value_text = fmt_pct(repa * 100)
                    m.status = "pass" if repa * 100 > 1.0 else ("warn" if repa * 100 > 0.7 else "fail")
                    m.detail = f"ROA dichiarato nel report: {fmt_pct(repa * 100)}"
    return res

_show_v9base = BuffettApp._show_result
def _show_v9(self, res):
    _show_v9base(self, res)
    try:
        self.nb.tab(self.tabs["quant"], text=("20  CRITERI QUANTITATIVI (BANKING)" if res["D"].get("_bank") else "CRITERI QUANTITATIVI"))
    except Exception:
        pass
BuffettApp._show_result = _show_v9

_show_v10base = BuffettApp._show_result
def _show_v10(self, res):
    _show_v10base(self, res)
    try:
        self.nb.tab(self.tabs["quant"],
                    text=("20  CRITERI QUANTITATIVI (BANKING)" if res["D"].get("_bank") else "CRITERI QUANTITATIVI"))
    except Exception:
        pass
BuffettApp._show_result = _show_v10

_show_v11base = BuffettApp._show_result
def _show_v11(self, res):
    _show_v11base(self, res)
    try:
        if res["D"].get("_bank"):
            n_vis = sum(1 for m in [m for m in res["quant"] if not (res["D"].get("_bank") and m.status == "nd")] if not getattr(m, "hidden", False) and m.status != "nd")
            self.nb.tab(self.tabs["quant"], text=f"{n_vis}  CRITERI QUANTITATIVI (BANKING)")
        else:
            self.nb.tab(self.tabs["quant"], text="CRITERI QUANTITATIVI")
    except Exception:
        pass
BuffettApp._show_result = _show_v11

def _yf_price_v11(company, text):
    tm = globals().get("_TICKER_MAP") or {
        "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
        "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK"}
    tick = None
    low = (company or "").lower()
    for k, vv in tm.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\s+([A-Z]{1,5})\b|\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if not tick: return None
    try:
        import yfinance as yf
        t = yf.Ticker(tick)
        for attr in ("fast_info", "info"):
            try:
                info = getattr(t, attr)
                for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                    try:
                        vv = info[k]
                        if vv: return float(vv)
                    except Exception: pass
            except Exception: pass
    except Exception:
        pass
    return None

_analyze_v11base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v11base(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p = _yf_price_v11(res.get("company", ""), text)
        if p is None:
            for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
                try:
                    t2 = open(pp, encoding="utf-8").read()
                    mm = re.search(r"([\d.,]+)", t2)
                    if mm:
                        snum = mm.group(1)
                        p = float(snum.replace(".", "").replace(",", ".")) if "," in snum else float(snum)
                        break
                except Exception: pass
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            res["log"].append(("ok", f"price         = {p:.2f} (mercato)"))
    return res

_TICKER_MAP_V12 = {
    "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
    "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK", "jpmorgan": "JPM",
}

def _fetch_price_v12(company, text):
    tick = None
    low = (company or "").lower()
    for k, vv in _TICKER_MAP_V12.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\s+([A-Z]{1,5})\b|\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if not tick: return None
    try:
        import yfinance as yf
        t = yf.Ticker(tick)
        for attr in ("fast_info", "info"):
            try:
                info = getattr(t, attr)
                for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                    try:
                        vv = info[k]
                        if vv: return float(vv)
                    except Exception: pass
            except Exception: pass
    except Exception:
        pass
    return None

_analyze_v12base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v12base(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p = _fetch_price_v12(res.get("company", ""), text)
        if p is None:
            for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
                try:
                    t2 = open(pp, encoding="utf-8").read()
                    mm = re.search(r"([\d.,]+)", t2)
                    if mm:
                        sn = mm.group(1)
                        p = float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)
                        break
                except Exception: pass
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            res["log"].append(("ok", f"price         = {p:.2f} (mercato)"))
    return res

# V13 price
_TICKER_V13 = {"unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
               "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK"}

def _price_v13(company, text):
    tick = None
    low = (company or "").lower()
    for k, vv in _TICKER_V13.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\s+([A-Z]{1,5})\b|\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if not tick: return None, "nessun ticker riconosciuto"
    try:
        import yfinance as yf
    except Exception:
        return None, "yfinance non installato"
    try:
        t = yf.Ticker(tick)
        for attr in ("fast_info", "info"):
            try:
                info = getattr(t, attr)
                for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                    try:
                        vv = info[k]
                        if vv: return float(vv), tick
                    except Exception: pass
            except Exception: pass
    except Exception:
        pass
    return None, f"ticker {tick} senza prezzo (rete assente?)"

_analyze_v13base = analyze_document
def analyze_document(path, cb=None):
    res = _analyze_v13base(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p, info = _price_v13(res.get("company", ""), text)
        if p is None:
            for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
                try:
                    t2 = open(pp, encoding="utf-8").read()
                    mm = re.search(r"([\d.,]+)", t2)
                    if mm:
                        sn = mm.group(1)
                        p = float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)
                        info = "prezzo.txt"; break
                except Exception: pass
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            res["log"].append(("ok", f"price         = {p:.2f} ({info})"))
        else:
            res["log"].append(("warn", f"price         = N/D ({info}) — crea prezzo.txt con es. 52,4"))
    return res

# ============================ V14 runtime-fix ============================
_TICKER_V14 = {"unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "bank of alexandria": "ALEX.CA",
               "apple": "AAPL", "microsoft": "MSFT", "blackrock": "BLK", "jpmorgan": "JPM"}

def _price_v14(company, text):
    tick = None
    low = (company or "").lower()
    for k, vv in _TICKER_V14.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\\s+([A-Z]{1,5})\\b|\\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\\s]+([A-Z0-9.\\-]{1,6})\\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if not tick: return None, "nessun ticker riconosciuto"
    try:
        import yfinance as yf
    except Exception:
        return None, "yfinance non installato"
    try:
        t = yf.Ticker(tick)
        for attr in ("fast_info", "info"):
            try:
                info = getattr(t, attr)
                for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                    try:
                        vv = info[k]
                        if vv: return float(vv), tick
                    except Exception: pass
            except Exception: pass
    except Exception:
        pass
    return None, f"ticker {tick} senza prezzo (rete assente?)"

_an_v14base = analyze_document
def analyze_document(path, cb=None):
    res = _an_v14base(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p, info = _price_v14(res.get("company", ""), text)
        if p is None:
            for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
                try:
                    t2 = open(pp, encoding="utf-8").read()
                    mm = re.search(r"([\\d.,]+)", t2)
                    if mm:
                        sn = mm.group(1)
                        p = float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)
                        info = "prezzo.txt"; break
                except Exception: pass
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            res["log"].append(("ok", f"price         = {p:.2f} ({info})"))
        else:
            res["log"].append(("warn", f"price         = N/D ({info}) — crea prezzo.txt con es. 52,4"))
    return res

_bqt_v14base = BuffettApp._build_quant_tab
def _bqt_v14(self, res):
    _bqt_v14base(self, res)
    try:
        f = self.tabs["quant"]
        bank = bool(res["D"].get("_bank"))
        for w in list(f.winfo_children()):
            try:
                t = w.cget("text")
            except Exception:
                continue
            if isinstance(t, str) and t.startswith("I 40 CRITERI"):
                w.configure(text="I 20 CRITERI QUANTITATIVI — CRITERI BANCARI" if bank else t)
        if bank:
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    txts = []
                    stack = [w]
                    while stack:
                        x = stack.pop()
                        for c in x.winfo_children():
                            try: txts.append(str(c.cget("text")))
                            except Exception: pass
                            stack.append(c)
                    if any("N/D" in t for t in txts):
                        w.destroy()
    except Exception:
        pass
BuffettApp._build_quant_tab = _bqt_v14

_sr_v14base = BuffettApp._show_result
def _sr_v14(self, res):
    _sr_v14base(self, res)
    try:
        self.nb.tab(self.tabs["quant"],
                    text=("20  CRITERI QUANTITATIVI" if res["D"].get("_bank") else "CRITERI QUANTITATIVI"))
    except Exception:
        pass
BuffettApp._show_result = _sr_v14

# ============================ V15 finale ============================
_PREZZI_STATICI = {"unicredit": 52.0, "intesa sanpaolo": 4.6, "bank of alexandria": 25.0,
                   "apple": 230.0, "microsoft": 500.0, "blackrock": 1000.0}

def _price_v15(company, text):
    tick = None
    low = (company or "").lower()
    mappa = globals().get("_TICKER_V14") or globals().get("_TICKER_MAP") or {
        "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "apple": "AAPL",
        "microsoft": "MSFT", "blackrock": "BLK"}
    for k, vv in mappa.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\s+([A-Z]{1,5})\b|\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\s]+([A-Z0-9.\-]{1,6})\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if tick:
        try:
            import yfinance as yf
            t = yf.Ticker(tick)
            for attr in ("fast_info", "info"):
                try:
                    info = getattr(t, attr)
                    for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                        try:
                            vv = info[k]
                            if vv: return float(vv), f"mercato {tick}"
                        except Exception: pass
                except Exception: pass
        except Exception:
            pass
    for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
        try:
            t2 = open(pp, encoding="utf-8").read()
            mm = re.search(r"([\d.,]+)", t2)
            if mm:
                sn = mm.group(1)
                return (float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)), "prezzo.txt"
        except Exception: pass
    for k, vv in _PREZZI_STATICI.items():
        if k in low: return vv, f"STATICO {vv} (aggiorna con prezzo.txt)"
    return None, "nessuna fonte prezzo"

_an_v15base = analyze_document
def analyze_document(path, cb=None):
    res = _an_v15base(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p, info = _price_v15(res.get("company", ""), text)
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            lvl = "warn" if "STATICO" in info else "ok"
            res["log"].append((lvl, f"price         = {p:.2f} ({info})"))
    return res

_bqt_v15base = BuffettApp._build_quant_tab
def _bqt_v15(self, res):
    _bqt_v15base(self, res)
    try:
        f = self.tabs["quant"]
        bank = bool(res["D"].get("_bank"))
        for w in list(f.winfo_children()):
            try: t = str(w.cget("text"))
            except Exception: continue
            if "40 CRITERI" in t:
                w.configure(text="I 20 CRITERI QUANTITATIVI — CRITERI BANCARI" if bank else t)
        if bank:
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    txts, stack = [], [w]
                    while stack:
                        x = stack.pop()
                        for c in x.winfo_children():
                            try: txts.append(str(c.cget("text")))
                            except Exception: pass
                            stack.append(c)
                    if any("N/D" in t for t in txts):
                        w.destroy()
    except Exception: pass
BuffettApp._build_quant_tab = _bqt_v15

_sr_v15base = BuffettApp._show_result
def _sr_v15(self, res):
    _sr_v15base(self, res)
    try:
        self.nb.tab(self.tabs["quant"],
                    text=("20  CRITERI QUANTITATIVI" if res["D"].get("_bank") else "40  CRITERI QUANTITATIVI"))
    except Exception: pass
BuffettApp._show_result = _sr_v15

# ============================ V16 velocità+log ============================
import time as _time16

def _structured_pdf_tables(path):
    """Versione con budget di tempo: max ~45s, mai più blocchi di mezz'ora."""
    try:
        import pdfplumber
    except Exception:
        return ""
    try:
        from pypdf import PdfReader as _R
    except Exception:
        try:
            from PyPDF2 import PdfReader as _R
        except Exception:
            return ""
    heads = globals().get("_STMT_HEADS") or ["stato patrimoniale", "conto economico",
        "consolidated balance sheets", "consolidated statements of operations",
        "statements of cash flows", "balance sheets", "income statements"]
    t0 = _time16.time()
    try:
        rdr = _R(path)
        n = len(rdr.pages)
        print(f"[V16] PDF: {n} pagine, scansione intestazioni (max 250)...")
        idxs = []
        for i in range(min(n, 250)):
            if _time16.time() - t0 > 20:
                print("[V16] budget 20s superato nella scansione: stop"); break
            t = (rdr.pages[i].extract_text() or "").lower()
            if any(h in t for h in heads):
                idxs.append(i)
            if len(idxs) >= 12: break
        if not idxs:
            print("[V16] nessuna pagina prospetti trovata: salto tabelle")
            return ""
        print(f"[V16] pagine prospetti: {idxs}")
        out = []
        with pdfplumber.open(path) as pdf:
            for i in idxs:
                if _time16.time() - t0 > 45:
                    print("[V16] budget 45s superato: stop tabelle"); break
                try:
                    tabs = pdf.pages[i].extract_tables()
                except Exception:
                    continue
                for tb in tabs or []:
                    for row in tb:
                        cells = [str(c).replace("\n", " ").strip() for c in row if c and str(c).strip()]
                        if len(cells) >= 2 and any(ch.isdigit() for ch in cells[-1]):
                            out.append(" | ".join(cells))
        print(f"[V16] tabelle: {len(out)} righe in {_time16.time()-t0:.1f}s")
        return "\n".join(out)
    except Exception as e:
        print(f"[V16] errore tabelle: {e}")
        return ""

def _yf_fetch_timeout(tick, secs=8):
    import threading
    box = {}
    def worker():
        try:
            import yfinance as yf
            t = yf.Ticker(tick)
            for attr in ("fast_info", "info"):
                try:
                    info = getattr(t, attr)
                    for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                        try:
                            vv = info[k]
                            if vv:
                                box["p"] = float(vv); return
                        except Exception: pass
                except Exception: pass
        except Exception: pass
    th = threading.Thread(target=worker, daemon=True)
    th.start(); th.join(secs)
    if th.is_alive():
        print(f"[V16] yfinance: timeout {secs}s su {tick} (rete assente?)")
    return box.get("p")

_price_v15_orig = globals().get("_price_v15")
def _price_v15(company, text):
    tick = None
    low = (company or "").lower()
    mappa = globals().get("_TICKER_V14") or globals().get("_TICKER_MAP") or {
        "unicredit": "UCG.MI", "intesa sanpaolo": "ISP.MI", "apple": "AAPL",
        "microsoft": "MSFT", "blackrock": "BLK"}
    for k, vv in mappa.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = re.search(r"under the symbol\\s+([A-Z]{1,5})\\b|\\b(?:BIT|MIL|NYSE|NASDAQ|LSE|EPA)[:\\s]+([A-Z0-9.\\-]{1,6})\\b", text)
        if m: tick = (m.group(1) or m.group(2))
    if tick:
        p = _yf_fetch_timeout(tick, 8)
        if p: return p, f"mercato {tick}"
    for pp in (os.path.join(BASE_DIR, "prezzo.txt"), os.path.join(OUTPUT_DIR, "prezzo.txt")):
        try:
            t2 = open(pp, encoding="utf-8").read()
            mm = re.search(r"([\\d.,]+)", t2)
            if mm:
                sn = mm.group(1)
                return (float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)), "prezzo.txt"
        except Exception: pass
    for k, vv in (globals().get("_PREZZI_STATICI") or {}).items():
        if k in low: return vv, f"STATICO {vv} (aggiorna con prezzo.txt)"
    return None, "nessuna fonte prezzo"

_an_v16base = analyze_document
def analyze_document(path, cb=None):
    print(f"[V16] >>> analisi: {os.path.basename(path)}")
    t0 = _time16.time()
    res = _an_v16base(path, cb)
    print(f"[V16] <<< completata in {_time16.time()-t0:.1f}s · score {res['scores']['final']:.1f}")
    return res

# ============================ V17 anti-blocco ============================
_ext_v17base = extract_text_from_file
def extract_text_from_file(path, *a, **k):
    import threading, time as _t17
    box = {}
    def w():
        try:
            box["t"] = _ext_v17base(path, *a, **k)
        except Exception as e:
            box["e"] = e
    th = threading.Thread(target=w, daemon=True)
    t0 = _t17.time()
    th.start(); th.join(120)
    if th.is_alive():
        print("[V17] estrazione > 120s: PDF troppo pesante/scannerizzato -> interrotta")
        raise ValueError("Estrazione troppo lunga (>120s): usa un PDF testuale o riduci le pagine")
    print(f"[V17] estrazione completata in {_t17.time()-t0:.1f}s")
    if "e" in box: raise box["e"]
    return box["t"]

# ============================ V18 estrazione budget ============================
_ext_v18base = extract_text_from_file
def extract_text_from_file(path, *a, **k):
    try:
        with open(path, "rb") as f: head = f.read(5)
    except Exception:
        head = b""
    ext = os.path.splitext(path)[1].lower()
    if not (head == b"%PDF-" or ext == ".pdf"):
        return _ext_v18base(path, *a, **k)
    import time as _t18, io as _io18
    try:
        from pypdf import PdfReader
    except Exception:
        from PyPDF2 import PdfReader
    t0 = _t18.time()
    with open(path, "rb") as f: data = f.read()
    r = PdfReader(_io18.BytesIO(data))
    n = len(r.pages)
    print(f"[V18] PDF: {n} pagine, estrazione testo (budget 90s)...")
    chunks = []
    for i, p in enumerate(r.pages):
        if _t18.time() - t0 > 90:
            print(f"[V18] budget 90s raggiunto a pagina {i}/{n}: continuo con il testo raccolto")
            break
        if i and i % 200 == 0:
            print(f"[V18]   ...pagina {i}/{n}")
        try:
            chunks.append(p.extract_text() or "")
        except Exception:
            chunks.append("")
    text = "\n".join(chunks)
    if len(text.strip()) < 200:
        print("[V18] testo quasi assente: PDF scannerizzato? OCR limitato (prime 8 pagine)")
        try:
            from pdf2image import convert_from_path
            import pytesseract
            imgs = convert_from_path(path, dpi=120, first_page=1, last_page=min(8, n))
            text = (text + "\n" + "\n".join(pytesseract.image_to_string(im, lang="eng+ita") for im in imgs)).strip()
        except Exception as e:
            print(f"[V18] OCR non disponibile: {e}")
    if text.strip():
        try:
            extra = _structured_pdf_tables(path)
            if extra:
                text += "\n" + extra
        except Exception:
            pass
    if not text.strip():
        raise ValueError("PDF senza testo estraibile (scannerizzato) e OCR non disponibile")
    print(f"[V18] estrazione OK in {_t18.time()-t0:.1f}s · {len(text)} caratteri")
    return text

# ============================ V19 banche-fix ============================
_an_v19base = analyze_document
def analyze_document(path, cb=None):
    res = _an_v19base(path, cb)
    D = res["D"]
    if D.get("_bank") and (D.get("price") is None or D.get("price") <= 1):
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        try:
            p, info = _price_v15(res.get("company", ""), text)
        except Exception:
            p, info = None, "prezzo non disponibile"
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            res["log"].append(("ok", f"price         = {p:.2f} ({info})"))
    return res

_sr_v19base = BuffettApp._show_result
def _sr_v19(self, res):
    _sr_v19base(self, res)
    try:
        bank = bool(res["D"].get("_bank"))
        self.nb.tab(self.tabs["quant"],
                    text=("20  CRITERI QUANTITATIVI" if bank else "40  CRITERI QUANTITATIVI"))
        f = self.tabs["quant"]
        for w in list(f.winfo_children()):
            try: t = str(w.cget("text"))
            except Exception: continue
            if "40 CRITERI" in t:
                w.configure(text="I 20 CRITERI QUANTITATIVI — CRITERI BANCARI" if bank else t)
        if bank:
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    txts, stack = [], [w]
                    while stack:
                        x = stack.pop()
                        for c in x.winfo_children():
                            try: txts.append(str(c.cget("text")))
                            except Exception: pass
                            stack.append(c)
                    if any("N/D" in t for t in txts):
                        w.destroy()
    except Exception: pass
BuffettApp._show_result = _sr_v19

# ============================ V20 coherence+UI ============================
_sr_v20base = BuffettApp._show_result
def _sr_v20(self, res):
    _sr_v20base(self, res)
    bank = bool(res["D"].get("_bank"))
    try:
        for t in self.nb.tabs():
            try:
                txt = str(self.nb.tab(t, "text"))
            except Exception:
                continue
            if "CRITERI QUANTITATIVI" in txt or "CRITERI BANCARI" in txt:
                self.nb.tab(t, text=("20  CRITERI BANCARI" if bank else "40  CRITERI QUANTITATIVI"))
    except Exception:
        pass
    try:
        f = self.tabs["quant"]
        for w in list(f.winfo_children()):
            try:
                t2 = str(w.cget("text"))
            except Exception:
                continue
            if "40 CRITERI" in t2:
                w.configure(text="I 20 CRITERI BANCARI — METRICHE DI SETTORE + VALUTAZIONE" if bank else t2)
    except Exception:
        pass
    if bank:
        try:
            f = self.tabs["quant"]
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    txts, stack = [], [w]
                    while stack:
                        x = stack.pop()
                        for c in x.winfo_children():
                            try:
                                txts.append(str(c.cget("text")))
                            except Exception:
                                pass
                            stack.append(c)
                    if any("N/D" in t for t in txts):
                        w.destroy()
        except Exception:
            pass
BuffettApp._show_result = _sr_v20

# ============================ V21 chiusura ============================
_extfin21 = extract_financials
def extract_financials(text, *a, **k):
    D, S, log = _extfin21(text, *a, **k)
    try:
        import json as _j21
        ov = _j21.load(open(os.path.join(BASE_DIR, "patterns_custom.json"), encoding="utf-8")).get("overrides", {})
    except Exception:
        ov = {}
    if D.get("_bank") and ov.get("eps") and D.get("net_income"):
        D["eps_reported"] = float(ov["eps"]); D["shares"] = D["net_income"] / float(ov["eps"])
        log.append(("ok", f"eps           = {ov['eps']} (override)"))
    if D.get("_bank") and ov.get("prezzo"):
        D["price"] = float(ov["prezzo"])
        if D.get("shares"): D["market_cap"] = D["price"] * D["shares"]
        log.append(("ok", f"price         = {ov['prezzo']} (override)"))
    return D, S, log

_sr21b = BuffettApp._show_result
def _sr_v21(self, res):
    _sr21b(self, res)
    bank = bool(res["D"].get("_bank"))
    try:
        for t in self.nb.tabs():
            tt = str(self.nb.tab(t, "text"))
            if "QUANTITATIVI" in tt or "BANCARI" in tt:
                self.nb.tab(t, text=("20  CRITERI BANCARI" if bank else "40  CRITERI QUANTITATIVI"))
    except Exception: pass
    try:
        f = self.tabs["quant"]
        for w in list(f.winfo_children()):
            try: t2 = str(w.cget("text"))
            except Exception: continue
            if "40 CRITERI" in t2: w.configure(text="I 20 CRITERI BANCARI")
        if bank:
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    bad, st = False, [w]
                    while st:
                        x = st.pop()
                        for c in x.winfo_children():
                            try: bad = bad or ("N/D" in str(c.cget("text")))
                            except Exception: pass
                            st.append(c)
                    if bad: w.destroy()
    except Exception: pass
BuffettApp._show_result = _sr_v21

# ============================ V22 finale ============================
_ef22b = extract_financials
def extract_financials(text, *a, **k):
    D, S, log = _ef22b(text, *a, **k)
    if D.get("_bank") and D.get("roe_reported") and D.get("net_income"):
        eq_imp = D["net_income"] / D["roe_reported"]
        if D.get("equity") is None or abs(D["equity"] - eq_imp) > 0.5 * eq_imp:
            D["equity"] = eq_imp
            log.append(("ok", f"equity        = {fmt(eq_imp,0)} (implicito da RoTE dichiarato)"))
    return D, S, log

_an22b = analyze_document
def analyze_document(path, cb=None):
    res = _an22b(path, cb)
    if res["D"].get("_bank"):
        for m in res["quant"]:
            if m.code == "Q10":
                m.status = "nd"; m.hidden = True
        res["scores"] = compute_scores(res["quant"], res["qual"])
    return res

_sr22b = BuffettApp._show_result
def _sr_v22(self, res):
    _sr22b(self, res)
    if res["D"].get("_bank"):
        try:
            f = self.tabs["quant"]
            for w in list(f.winfo_children()):
                if str(w.winfo_class()) == "Frame":
                    blob, st = [], [w]
                    while st:
                        x = st.pop()
                        for c in x.winfo_children():
                            try: blob.append(str(c.cget("text")))
                            except Exception: pass
                            st.append(c)
                    txt = " ".join(blob)
                    if ("N/D" in txt) or ("Test del $1" in txt):
                        w.destroy()
        except Exception: pass
BuffettApp._show_result = _sr_v22

# ============================ V25 anchor-free ============================
_an25b = analyze_document
def analyze_document(path, cb=None):
    res = _an25b(path, cb)
    D = res["D"]
    if D.get("_bank") and D.get("roe_reported") and D.get("net_income") and D.get("shares"):
        eq = D["net_income"] / D["roe_reported"]
        bvps = eq / D["shares"]
        tbv = eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)
        tbvps = tbv / D["shares"] if D["shares"] else None
        pr = D.get("price")
        for m in res["quant"]:
            if m.code == "Q34" and bvps:
                m.value = pr / bvps if pr else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(pr,2)} / BVPS {fmt(bvps,2)}"
                if m.value: m.status = "pass" if m.value < 3 else ("warn" if m.value < 5 else "fail")
            if m.code == "B3" and tbvps:
                m.value = pr / tbvps if pr else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(pr,2)} / TBVPS {fmt(tbvps,2)}"
                if m.value: m.status = "pass" if m.value < 1.5 else ("warn" if m.value < 2 else "fail")
        print(f"[V25] BVPS={bvps:.2f} TBVPS={tbvps:.2f} -> P/BV e P/TBV riscritti")
    return res

_sr25b = BuffettApp._show_result
def _sr_v25(self, res):
    _sr25b(self, res)
    bank = bool(res["D"].get("_bank"))
    hit = 0
    try:
        for t in self.nb.tabs():
            tt = str(self.nb.tab(t, "text"))
            if "CRITERI QUANTITATIVI" in tt or "CRITERI BANCARI" in tt:
                self.nb.tab(t, text=("20  CRITERI BANCARI" if bank else "40  CRITERI QUANTITATIVI")); hit += 1
    except Exception: pass
    try:
        f = self.tabs["quant"]
        st = [f]
        while st:
            x = st.pop()
            for c in x.winfo_children():
                try:
                    t2 = str(c.cget("text"))
                    if "40 CRITERI" in t2:
                        c.configure(text="I 20 CRITERI BANCARI — METRICHE DI SETTORE + VALUTAZIONE" if bank else t2); hit += 1
                except Exception: pass
                st.append(c)
    except Exception: pass
    if hit: print(f"[V25] UI rinominata ({hit} elementi)")
BuffettApp._show_result = _sr_v25

# ============================ V26 definitivo ============================
_an26b = analyze_document
def analyze_document(path, cb=None):
    res = _an26b(path, cb)
    D = res["D"]
    if D.get("_bank") and D.get("roe_reported") and D.get("net_income") and D.get("shares"):
        eq = D["net_income"] / D["roe_reported"]
        bvps = eq / D["shares"]
        tbvps = (eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)) / D["shares"]
        for m in res["quant"]:
            if m.code == "Q34":
                m.value = (D["price"] / bvps) if D.get("price") else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(D.get('price'),2)} / BVPS {fmt(bvps,2)}"
                m.status = "pass" if m.value and m.value < 3 else ("warn" if m.value and m.value < 5 else "fail")
            if m.code == "B3":
                m.value = (D["price"] / tbvps) if (D.get("price") and tbvps) else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(D.get('price'),2)} / TBVPS {fmt(tbvps,2)}"
                m.status = "pass" if m.value and m.value < 1.5 else ("warn" if m.value and m.value < 2.0 else "fail")
    return res

_sr26b = BuffettApp._show_result
def _sr_v26(self, res):
    _sr26b(self, res)
    bank = bool(res["D"].get("_bank"))
    def walk(w):
        try:
            for c in w.winfo_children():
                try:
                    t = str(c.cget("text"))
                    if "40 CRITERI QUANTITATIVI" in t:
                        c.configure(text=("I 20 CRITERI BANCARI — METRICHE DI SETTORE + VALUTAZIONE" if bank else t))
                except Exception:
                    pass
                walk(c)
        except Exception:
            pass
    walk(self.root)
    try:
        for t in self.nb.tabs():
            if "CRITERI QUANTITATIVI" in str(self.nb.tab(t, "text")):
                self.nb.tab(t, text=("20  CRITERI BANCARI" if bank else "40  CRITERI QUANTITATIVI"))
    except Exception:
        pass
BuffettApp._show_result = _sr_v26

# ============================ V27 chiusura ============================
_an27b = analyze_document
def analyze_document(path, cb=None):
    res = _an27b(path, cb)
    D = res["D"]
    if D.get("_bank") and D.get("roe_reported") and D.get("net_income") and D.get("shares"):
        eq = D["net_income"] / D["roe_reported"]
        bvps = eq / D["shares"]
        tbvps = (eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)) / D["shares"]
        for m in res["quant"]:
            if m.code == "Q34":
                m.value = (D["price"] / bvps) if D.get("price") else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(D.get('price'),2)} / BVPS {fmt(bvps,2)}"
                m.status = "pass" if m.value and m.value < 3 else ("warn" if m.value and m.value < 5 else "fail")
            if m.code == "B3":
                m.value = (D["price"] / tbvps) if (D.get("price") and tbvps) else None
                m.value_text = fmt(m.value, 2) if m.value else "N/D"
                m.detail = f"Prezzo {fmt(D.get('price'),2)} / TBVPS {fmt(tbvps,2)}"
                m.status = "pass" if m.value and m.value < 1.5 else ("warn" if m.value and m.value < 2.0 else "fail")
    return res

_sr27b = BuffettApp._show_result
def _sr_v27(self, res):
    _sr27b(self, res)
    bank = bool(res["D"].get("_bank"))
    def walk(w):
        try:
            for c in w.winfo_children():
                try:
                    if "40 CRITERI QUANTITATIVI" in str(c.cget("text")):
                        c.configure(text=("I 20 CRITERI BANCARI — METRICHE DI SETTORE + VALUTAZIONE" if bank else str(c.cget("text"))))
                except Exception: pass
                walk(c)
        except Exception: pass
    walk(self.root)
    try:
        for t in self.nb.tabs():
            if "CRITERI QUANTITATIVI" in str(self.nb.tab(t, "text")):
                self.nb.tab(t, text=("20  CRITERI BANCARI" if bank else "40  CRITERI QUANTITATIVI"))
    except Exception: pass
BuffettApp._show_result = _sr_v27

# V28
_bq28 = build_quant
def build_quant(D, S, dcf):
    Q = _bq28(D, S, dcf)
    if D.get("_bank") and D.get("roe_reported") and D.get("net_income") and D.get("shares"):
        eq = D["net_income"] / D["roe_reported"]
        bvps = eq / D["shares"]
        tbvps = (eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)) / D["shares"]
        for m in Q:
            if m.code == "Q34" and D.get("price"):
                m.value = D["price"] / bvps; m.value_text = fmt(m.value, 2)
                m.detail = f"Prezzo {fmt(D['price'],2)} / BVPS {fmt(bvps,2)}"
                m.status = "pass" if m.value < 3 else ("warn" if m.value < 5 else "fail")
            if m.code == "B3" and D.get("price") and tbvps:
                m.value = D["price"] / tbvps; m.value_text = fmt(m.value, 2)
                m.detail = f"Prezzo {fmt(D['price'],2)} / TBVPS {fmt(tbvps,2)}"
                m.status = "pass" if m.value < 1.5 else ("warn" if m.value < 2.0 else "fail")
    return Q

_eh28 = export_html
def export_html(res):
    path = _eh28(res)
    try:
        h = open(path, encoding="utf-8").read()
        if res["D"].get("_bank"):
            h2 = h.replace("40 CRITERI QUANTITATIVI", "20 CRITERI BANCARI")
            if h2 != h: open(path, "w", encoding="utf-8").write(h2)
    except Exception:
        pass
    return path

# V29
_an29b = analyze_document
def analyze_document(path, cb=None):
    res = _an29b(path, cb)
    D = res["D"]
    roe = D.get("roe_reported") or D.get("rote")
    if D.get("_bank") and roe and D.get("net_income") and D.get("shares") and D.get("price"):
        eq = D["net_income"] / roe
        bvps = eq / D["shares"]
        tbvps = (eq - (D.get("goodwill") or 0) - (D.get("intangibles") or 0)) / D["shares"]
        for m in res["quant"]:
            if m.code == "Q34":
                m.value = D["price"] / bvps; m.value_text = fmt(m.value, 2)
                m.detail = f"Prezzo {fmt(D['price'],2)} / BVPS {fmt(bvps,2)}"
                m.status = "pass" if m.value < 3 else ("warn" if m.value < 5 else "fail")
            if m.code == "B3":
                m.value = D["price"] / tbvps; m.value_text = fmt(m.value, 2)
                m.detail = f"Prezzo {fmt(D['price'],2)} / TBVPS {fmt(tbvps,2)}"
                m.status = "pass" if m.value < 1.5 else ("warn" if m.value < 2 else "fail")
        print(f"[V29] BVPS={bvps:.2f} TBVPS={tbvps:.2f} -> Q34/B3 riscritti")
    return res

# V31
_sr31b = BuffettApp._show_result
def _sr_v31(self, res):
    _sr31b(self, res)
    D = res["D"]
    if not D.get("_bank"):
        return
    roe = D.get("roe_reported") or D.get("rote")
    ni, ta = D.get("net_income"), D.get("total_assets")
    roa = D.get("roa_reported") or (ni / ta if ni and ta else None)
    cet1 = D.get("cet1")
    repl = {
        "ROE": ("ROE", fmt_pct(roe * 100) if roe else None),
        "ROIC": ("ROA (BANCA)", fmt_pct(roa * 100) if roa else None),
        "DEBT/EQUITY": ("CET1 RATIO", fmt_pct(cet1 * 100) if cet1 else None),
    }
    def walk(w):
        try:
            kids = w.winfo_children()
        except Exception:
            return
        for c in kids:
            t = None
            try:
                t = str(c.cget("text")).strip().upper()
            except Exception:
                pass
            if t in repl:
                newlab, newval = repl[t]
                try:
                    c.configure(text=newlab)
                except Exception:
                    pass
                if newval:
                    for s2 in c.master.winfo_children():
                        if s2 is not c:
                            try:
                                s2.configure(text=newval)
                            except Exception:
                                pass
            walk(c)
    tgt = self.tabs.get("dash") or self.root
    walk(tgt)
    print("[V31] tile dashboard bancarie: ROE/ROA/CET1 popolate")
BuffettApp._show_result = _sr_v31

# V32
_sr32b = BuffettApp._show_result
def _sr_v32(self, res):
    _sr32b(self, res)
    if not res["D"].get("_bank"):
        return
    tgt = self.tabs.get("dash") or self.root
    def walk(w):
        try:
            kids = w.winfo_children()
        except Exception:
            return
        for c in kids:
            t = None
            try:
                t = str(c.cget("text")).strip().upper()
            except Exception:
                pass
            if t in ("CET1", "COST/INCOME"):
                try:
                    c.master.destroy()
                except Exception:
                    pass
            else:
                walk(c)
    walk(tgt)
    print("[V32] tile duplicate rimosse dalla dashboard bancaria")
BuffettApp._show_result = _sr_v32

# V33
_an33b = analyze_document
def analyze_document(path, cb=None):
    res = _an33b(path, cb)
    D = res["D"]
    if not D.get("ev") and D.get("market_cap"):
        D["ev"] = D["market_cap"] + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
        res["log"].append(("ok", f"ev            = {fmt(D['ev'],0)} (mcap+debt-cassa)"))
    if not D.get("oe") and D.get("fcf"):
        D["oe"] = D["fcf"]
    return res

# V34 price
_PRICE_CACHE = {}

def _fetch_price_any(company, text):
    """Cerca ticker nel testo o dal nome, scarica da yfinance, fallback prezzo.txt"""
    import re as _r
    tm = globals().get("_TICKER_V14") or {}
    tick = None
    low = (company or "").lower()
    for k, vv in tm.items():
        if k in low: tick = vv; break
    if not tick and text:
        m = _r.search(r"under the symbol\s+([A-Z]{1,5})\b", text)
        if m: tick = m.group(1)
    if not tick and text:
        m = _r.search(r"\b(?:NASDAQ|NYSE|BIT|MIL)[:\s]+([A-Z]{1,5})\b", text)
        if m: tick = m.group(1)
    if tick and tick in _PRICE_CACHE:
        return _PRICE_CACHE[tick], f"cache {tick}"
    if tick:
        try:
            import yfinance as yf
            t = yf.Ticker(tick)
            for attr in ("fast_info", "info"):
                try:
                    info = getattr(t, attr)
                    for k in ("currentPrice", "regularMarketPrice", "previousClose"):
                        try:
                            vv = info[k]
                            if vv:
                                _PRICE_CACHE[tick] = float(vv)
                                return float(vv), f"{tick}"
                        except Exception: pass
                except Exception: pass
        except Exception: pass
    for pp in ("prezzo.txt",):
        try:
            txt = open(pp, encoding="utf-8").read()
            mm = _r.search(r"([\d.,]+)", txt)
            if mm:
                sn = mm.group(1)
                p = float(sn.replace(".", "").replace(",", ".")) if "," in sn else float(sn)
                return p, "prezzo.txt"
        except Exception: pass
    return None, "nessuna fonte"

_an34b = analyze_document
def analyze_document(path, cb=None):
    res = _an34b(path, cb)
    D = res["D"]
    if D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        p, info = _fetch_price_any(res.get("company", ""), text)
        if p:
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            if D.get("market_cap"):
                D["ev"] = D["market_cap"] + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
            res["log"].append(("ok", f"price         = {p:.2f} ({info})"))
    return res

def _apply_manual_price(self, price_val):
    """Input manuale: ricalcola tutte le metriche di mercato"""
    if not getattr(self, "result", None): return
    D = self.result["D"]
    D["price"] = price_val
    if D.get("shares"): D["market_cap"] = price_val * D["shares"]
    if D.get("market_cap"):
        D["ev"] = D["market_cap"] + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
    # ricalcola le card di valutazione
    eps = D.get("net_income") / D.get("shares") if D.get("shares") else None
    bvps = D.get("equity") / D.get("shares") if D.get("shares") and D.get("equity") else None
    fcf = D.get("fcf")
    for m in self.result["quant"]:
        c = m.code
        if c == "Q32" and eps:
            m.value = price_val / eps; m.value_text = fmt(m.value, 1)
            m.status = "pass" if m.value < 15 else ("warn" if m.value < 25 else "fail")
        if c == "Q34" and bvps:
            m.value = price_val / bvps; m.value_text = fmt(m.value, 2)
            m.status = "pass" if m.value < 3 else ("warn" if m.value < 5 else "fail")
            m.detail = f"Prezzo {fmt(price_val,2)} / BVPS {fmt(bvps,2)}"
        if c == "Q16" and eps:
            m.value = eps / price_val * 100; m.value_text = fmt_pct(m.value)
            m.status = "pass" if m.value > 7 else ("warn" if m.value > 5 else "fail")
        if c == "Q17" and fcf and D.get("ev"):
            m.value = fcf / D["ev"] * 100; m.value_text = fmt_pct(m.value)
        if c == "Q35" and D.get("ebit") and D.get("ev"):
            m.value = D["ev"] / D["ebit"]; m.value_text = fmt(m.value, 1)
        if c == "Q36" and fcf and D.get("ev"):
            m.value = D["ev"] / fcf; m.value_text = fmt(m.value, 1)
        if c == "Q09" and self.result.get("dcf", {}).get("iv"):
            iv = self.result["dcf"]["iv"]
            mcap = D.get("market_cap")
            if mcap:
                mos = (iv - mcap) / iv * 100
                m.value = mos; m.value_text = fmt_pct(mos)
                m.status = "pass" if mos >= 30 else ("warn" if mos >= 20 else "fail")
    self.result["scores"] = compute_scores(self.result["quant"], self.result["qual"])
    self._show_result(self.result)
BuffettApp._apply_manual_price = _apply_manual_price

_sr34b = BuffettApp._show_result
def _sr_v34(self, res):
    _sr34b(self, res)
    try:
        top = getattr(self, "top_bar", None) or self.root
        if not hasattr(self, "_price_entry"):
            import tkinter as tk
            f = tk.Frame(top)
            f.pack(side="top", fill="x", pady=2)
            tk.Label(f, text="Prezzo azione:").pack(side="left", padx=(8,4))
            self._price_entry = tk.Entry(f, width=10)
            self._price_entry.pack(side="left")
            if res["D"].get("price"):
                self._price_entry.insert(0, str(round(res["D"]["price"], 2)))
            def go():
                try:
                    v = float(self._price_entry.get().replace(",", "."))
                    self._apply_manual_price(v)
                except Exception: pass
            tk.Button(f, text="Applica prezzo", command=go).pack(side="left", padx=4)
    except Exception: pass
BuffettApp._show_result = _sr_v34

# V35 buyback+debt+price
_FALLBACK_PRICES = {"MSFT": 425.0, "AAPL": 230.0, "GOOGL": 170.0, "GOOG": 170.0,
                    "AMZN": 200.0, "NVDA": 130.0, "META": 600.0, "BRK-B": 500.0}

_an35b = analyze_document
def analyze_document(path, cb=None):
    res = _an35b(path, cb)
    D = res["D"]
    bank = bool(D.get("_bank"))

    # 3) prezzo fallback: se yfinance non risponde, uso tabella statica
    if not bank and D.get("price") is None:
        text = ""
        try:
            with open(os.path.join(OUTPUT_DIR, "_ultimo_testo_estratto.txt"), encoding="utf-8") as f:
                text = f.read()
        except Exception: pass
        import re as _r
        tick = None
        for m in _r.finditer(r"under the symbol\s+([A-Z]{1,5})\b", text):
            tick = m.group(1); break
        if not tick:
            for m in _r.finditer(r"\b(?:NASDAQ|NYSE)[:\s]+([A-Z]{1,5})\b", text):
                tick = m.group(1); break
        if tick and tick.upper() in _FALLBACK_PRICES:
            p = _FALLBACK_PRICES[tick.upper()]
            D["price"] = p
            if D.get("shares"): D["market_cap"] = p * D["shares"]
            if D.get("market_cap"):
                D["ev"] = D["market_cap"] + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
            res["log"].append(("warn", f"price         = {p:.2f} (fallback {tick} - aggiorna con prezzo.txt)"))

    if bank:
        return res

    # 1) ROE/ROIC: cap 100% + avviso buyback + ROIC su Capitale Operativo Netto
    ni = D.get("net_income"); eq = D.get("equity"); ta = D.get("total_assets")
    cassa = D.get("cassa") or 0; cl = D.get("current_liab") or 0
    if eq and ni and ni / eq > 1.0:
        # buyback massiccio: usa equity normalizzato (media 5 anni se disponibile)
        sn = D.get("equity_hist") or []
        if len(sn) >= 2:
            eq_avg = sum(sn) / len(sn)
            D["equity_normalized"] = eq_avg
        else:
            D["equity_normalized"] = ta * 0.35  # stima settore tech
        res["log"].append(("warn", "ROE distorto da buyback: uso equity normalizzato"))

    # ROIC su Capitale Operativo Netto
    nopat = (D.get("ebit") or 0) * (1 - (D.get("tax_rate") or 0.25))
    cap_op = (ta or 0) - cassa - (cl or 0)
    if cap_op > 0:
        D["roic_operativo"] = nopat / cap_op

    # 2) paradosso debito: Q04 prevale su Q05/Q39
    yrs = D.get("yrs_debt_payoff")
    if yrs is None and D.get("lt_debt") and ni:
        yrs = D["lt_debt"] / ni
        D["yrs_debt_payoff"] = yrs

    # riscrivo le card quantitative
    for m in res["quant"]:
        # ROE cap + buyback note
        if m.code == "Q03" and eq and ni:
            roe_raw = ni / eq * 100
            eq_n = D.get("equity_normalized")
            if roe_raw > 100 and eq_n:
                roe_n = ni / eq_n * 100
                m.value = roe_n; m.value_text = fmt_pct(roe_n)
                m.status = "pass" if roe_n > 15 else ("warn" if roe_n > 10 else "fail")
                m.detail = f"ROE contabile {roe_raw:.0f}% (buyback) -> normalizzato {roe_n:.1f}%"
        # ROIC operativo
        if m.code == "Q02" and D.get("roic_operativo"):
            v = D["roic_operativo"] * 100
            m.value = v; m.value_text = fmt_pct(v)
            m.status = "pass" if v > 15 else ("warn" if v > 10 else "fail")
            m.detail = "ROIC su Capitale Operativo Netto (esclusi buyback)"
        # paradosso debito
        if m.code in ("Q05", "Q39") and yrs is not None and yrs < 2:
            m.status = "warn"
            m.detail = f"Debito coperto in {yrs:.2f} anni: leva alta solo per buyback"

    # ricalcolo metriche di mercato con il prezzo (ora sempre presente)
    price = D.get("price"); eps = ni / D["shares"] if ni and D.get("shares") else None
    bvps = eq / D["shares"] if eq and D.get("shares") else None
    for m in res["quant"]:
        if m.code == "Q32" and eps and price:
            m.value = price / eps; m.value_text = fmt(m.value, 1)
            m.status = "pass" if m.value < 15 else ("warn" if m.value < 25 else "fail")
        if m.code == "Q34" and bvps and price:
            m.value = price / bvps; m.value_text = fmt(m.value, 2)
            m.status = "pass" if m.value < 3 else ("warn" if m.value < 5 else "fail")
            m.detail = f"Prezzo {fmt(price,2)} / BVPS {fmt(bvps,2)}"
        if m.code == "Q16" and eps and price:
            m.value = eps / price * 100; m.value_text = fmt_pct(m.value)
            m.status = "pass" if m.value > 7 else ("warn" if m.value > 5 else "fail")
        if m.code == "Q17" and D.get("fcf") and D.get("ev"):
            m.value = D["fcf"] / D["ev"] * 100; m.value_text = fmt_pct(m.value)
            m.status = "pass" if m.value > 7 else ("warn" if m.value > 5 else "fail")
        if m.code == "Q35" and D.get("ebit") and D.get("ev"):
            m.value = D["ev"] / D["ebit"]; m.value_text = fmt(m.value, 1)
        if m.code == "Q36" and D.get("fcf") and D.get("ev"):
            m.value = D["ev"] / D["fcf"]; m.value_text = fmt(m.value, 1)
        if m.code == "Q09" and D.get("market_cap") and res.get("dcf", {}).get("iv"):
            iv = res["dcf"]["iv"]
            mos = (iv - D["market_cap"]) / iv * 100
            m.value = mos; m.value_text = fmt_pct(mos)
            m.status = "pass" if mos >= 30 else ("warn" if mos >= 20 else "fail")

    res["scores"] = compute_scores(res["quant"], res["qual"])
    return res

# V36 ev-fix
_an36b = analyze_document
def analyze_document(path, cb=None):
    res = _an36b(path, cb)
    D = res["D"]
    # calcolo garantito market_cap e ev PRIMA delle metriche
    if D.get("price") and D.get("shares") and not D.get("market_cap"):
        D["market_cap"] = D["price"] * D["shares"]
        res["log"].append(("ok", f"market_cap    = {fmt(D['market_cap'],0)} (price×shares)"))
    if D.get("market_cap") and not D.get("ev"):
        D["ev"] = D["market_cap"] + (D.get("total_debt") or 0) - (D.get("cassa") or 0)
        res["log"].append(("ok", f"ev            = {fmt(D['ev'],0)} (mcap+debt-cash)"))
    # ricalcolo metriche EV-based se ev ora disponibile
    if D.get("ev") and not D.get("_bank"):
        fcf = D.get("fcf"); ebit = D.get("ebit")
        for m in res["quant"]:
            if m.code == "Q17" and fcf:
                m.value = fcf / D["ev"] * 100; m.value_text = fmt_pct(m.value)
                m.status = "pass" if m.value > 7 else ("warn" if m.value > 5 else "fail")
                m.detail = f"FCF {fmt(fcf,0)} / EV {fmt(D['ev'],0)}"
            if m.code == "Q35" and ebit:
                m.value = D["ev"] / ebit; m.value_text = fmt(m.value, 1)
                m.status = "pass" if m.value < 12 else ("warn" if m.value < 15 else "fail")
                m.detail = f"EV {fmt(D['ev'],0)} / EBIT {fmt(ebit,0)}"
            if m.code == "Q36" and fcf:
                m.value = D["ev"] / fcf; m.value_text = fmt(m.value, 1)
                m.status = "pass" if m.value < 15 else ("warn" if m.value < 20 else "fail")
                m.detail = f"EV {fmt(D['ev'],0)} / FCF {fmt(fcf,0)}"
    res["scores"] = compute_scores(res["quant"], res["qual"])
    return res

if __name__ == "__main__":
    main()